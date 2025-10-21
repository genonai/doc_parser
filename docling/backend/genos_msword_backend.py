import logging 
import re
from io import BytesIO
from pathlib import Path
from typing import Any, List, Optional, Union
from collections import defaultdict

# PIL ë¡œê¹… ë¹„í™œì„±í™” (FpxImagePlugin ì˜¤ë¥˜ ë©”ì‹œì§€ ë°©ì§€)
logging.getLogger('PIL').setLevel(logging.WARNING)
try:
    from wand.image import Image as WandImage
    from wand.exceptions import WandException
    WAND_AVAILABLE = True
except ImportError:
    WAND_AVAILABLE = False

from docling_core.types.doc.base import BoundingBox
from docling_core.types.doc.document import (
    DoclingDocument,
    DocumentOrigin,
    ImageRef,
    NodeItem,
    ProvenanceItem,
    TableCell,
    TableData,
)
from docling_core.types.doc.labels import DocItemLabel, GroupLabel
from docling_core.types.doc.document import Formatting
from docling_core.types.doc.base import Size
from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.xmlchemy import BaseOxmlElement
from docx.table import Table, _Cell
from docx.text.hyperlink import Hyperlink
from docx.text.paragraph import Paragraph
from docx.text.run import Run
import lxml.etree as etree
from PIL import Image, UnidentifiedImageError
from pydantic import AnyUrl
from typing_extensions import override

from docling.backend.abstract_backend import DeclarativeDocumentBackend
from docling.backend.docx.latex.omml import oMath2Latex
from docling.datamodel.base_models import InputFormat
from docling.datamodel.document import InputDocument

_log = logging.getLogger(__name__)


class GenosMsWordDocumentBackend(DeclarativeDocumentBackend):
    @override
    def __init__(
        self, in_doc: "InputDocument", path_or_stream: Union[BytesIO, Path]
    ) -> None:
        super().__init__(in_doc, path_or_stream)
        self.XML_KEY = (
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val"
        )
        self.xml_namespaces = {
            "w": "http://schemas.microsoft.com/office/word/2003/wordml"
        }
        # Word file:
        self.path_or_stream: Optional[Union[BytesIO, Path]] = path_or_stream
        self.valid: bool = False
        # Initialise the parents for the hierarchy
        self.max_levels: int = 10
        self.level_at_new_list: Optional[int] = None
        self.parents: dict[int, Optional[NodeItem]] = {}
        self.numbered_headers: dict[int, int] = {}
        self.equation_bookends: str = "<eq>{EQ}</eq>"
        # Track processed textbox elements to avoid duplication
        self.processed_textbox_elements: List[int] = []
        # Track content hash of processed paragraphs to avoid duplicate content
        self.processed_paragraph_content: List[str] = []
        # Track seen section texts for header detection (from hwpx_backend)
        self._seen_section_texts: set[str] = set()
        self.processed_table_elements: set[int] = set()

        for i in range(-1, self.max_levels):
            self.parents[i] = None

        self.level = 0
        self.listIter = 0
        self._seen_sectpr_ids = set()
        self.history: dict[str, Any] = {
            "names": [None],
            "levels": [None],
            "numids": [None],
            "indents": [None],
        }

        self.docx_obj = None
        try:
            # ğŸ”§ BytesIOëŠ” í¬ì¸í„°ë¥¼ ì²˜ìŒìœ¼ë¡œ
            if isinstance(self.path_or_stream, BytesIO):
                try:
                    self.path_or_stream.seek(0)
                except Exception:
                    pass
                self.docx_obj = Document(self.path_or_stream)

            elif isinstance(self.path_or_stream, Path):
                self.docx_obj = Document(str(self.path_or_stream))
            else:
                raise TypeError(f"Unsupported path_or_stream type: {type(self.path_or_stream)}")

            # âœ… ì—¬ê¸°ì„œë§Œ valid=True ë° package í• ë‹¹
            self.valid = True
            # âš ï¸ self.docx_objê°€ ìœ íš¨í•  ë•Œë§Œ ì ‘ê·¼
            self.package = self.docx_obj.part.package

        except Exception as e:
            # ë¡œë”© ì‹¤íŒ¨ ì‹œ ëª…í™•í•œ ë©”ì‹œì§€ + ì›ì¸ ìœ ì§€
            raise RuntimeError(
                f"GenosMsWordDocumentBackend could not load document with hash {self.document_hash}"
            ) from e
        try:
            if isinstance(self.path_or_stream, BytesIO):
                self.docx_obj = Document(self.path_or_stream)
            elif isinstance(self.path_or_stream, Path):
                self.docx_obj = Document(str(self.path_or_stream))

            self.valid = True
        except Exception as e:
            raise RuntimeError(
                f"GenosMsWordDocumentBackend could not load document with hash {self.document_hash}"
            ) from e

    @override
    def is_valid(self) -> bool:
        return self.valid

    @classmethod
    @override
    def supports_pagination(cls) -> bool:
        return False

    @override
    def unload(self):
        if isinstance(self.path_or_stream, BytesIO):
            self.path_or_stream.close()

        self.path_or_stream = None

    @classmethod
    @override
    def supported_formats(cls) -> set[InputFormat]:
        return {InputFormat.DOCX}

    @override
    def convert(self) -> DoclingDocument:
        """Parses the DOCX into a structured document model.

        Returns:
            The parsed document.
        """

        # binary_hash ëŠ” Uint64 ë¡œ ê¸°ëŒ€ë˜ë¯€ë¡œ, í•´ì‹œ ë¬¸ìì—´ì„ 64ë¹„íŠ¸ ì •ìˆ˜ë¡œ ì¶•ì†Œ
        try:
            bin_hash = int(self.document_hash, 16) & ((1 << 64) - 1)
        except Exception:
            bin_hash = 0

        origin = DocumentOrigin(
            filename=self.file.name or "file",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            binary_hash=bin_hash,
        )

        doc = DoclingDocument(name=self.file.stem or "file", origin=origin)

        if not self.is_valid():
            raise RuntimeError(
                f"Cannot convert doc with {self.document_hash} because the backend failed to init."
            )
        # ğŸ” document.xml(body)ë§Œ ì‹œì‘ì ìœ¼ë¡œ ìˆœì°¨ ì²˜ë¦¬
        assert self.docx_obj is not None
        body_el = self.docx_obj.element.body
        doc = self._walk_linear(
            body=body_el,
            docx_obj=self.docx_obj,
            doc=doc,
            owner_part=self.docx_obj.part,  # â† í˜„ì¬ XMLì˜ ì†Œìœ  íŒŒíŠ¸
        )

        doc.pages[1] =  doc.add_page(page_no=1, size=Size(width=595, height=842))
        return doc


    def _update_history(
        self,
        name: str,
        level: Optional[int],
        numid: Optional[int],
        ilevel: Optional[int],
    ):
        self.history["names"].append(name)
        self.history["levels"].append(level)
        self.history["numids"].append(numid)
        self.history["indents"].append(ilevel)

    def _prev_name(self) -> Optional[str]:
        return self.history["names"][-1]

    def _prev_level(self) -> Optional[int]:
        return self.history["levels"][-1]

    def _prev_numid(self) -> Optional[int]:
        return self.history["numids"][-1]

    def _prev_indent(self) -> Optional[int]:
        return self.history["indents"][-1]

    def _get_level(self) -> int:
        """Return the first None index."""
        for k, v in self.parents.items():
            if k >= 0 and v is None:
                return k
        return 0

    def _str_to_int(
        self, s: Optional[str], default: Optional[int] = 0
    ) -> Optional[int]:
        if s is None:
            return None
        try:
            return int(s)
        except ValueError:
            return default

    def _split_text_and_number(self, input_string: str) -> list[str]:
        match = re.match(r"(\D+)(\d+)$|^(\d+)(\D+)", input_string)
        if match:
            parts = list(filter(None, match.groups()))
            return parts
        else:
            return [input_string]

    def _get_numId_and_ilvl(
        self, paragraph: Paragraph
    ) -> tuple[Optional[int], Optional[int]]:
        # Access the XML element of the paragraph
        numPr = paragraph._element.find(
            ".//w:numPr", namespaces=paragraph._element.nsmap
        )

        if numPr is not None:
            # Get the numId element and extract the value
            numId_elem = numPr.find("w:numId", namespaces=paragraph._element.nsmap)
            ilvl_elem = numPr.find("w:ilvl", namespaces=paragraph._element.nsmap)
            numId = numId_elem.get(self.XML_KEY) if numId_elem is not None else None
            ilvl = ilvl_elem.get(self.XML_KEY) if ilvl_elem is not None else None

            return self._str_to_int(numId, None), self._str_to_int(ilvl, None)

        return None, None  # If the paragraph is not part of a list

    def _get_style_numId_and_ilvl(self, paragraph: Paragraph) -> tuple[Optional[int], Optional[int]]:
        """
        Try to resolve numId/ilvl from the paragraph's style definition if present.
        This helps when numbering is attached via style rather than inline paragraph properties.
        """
        try:
            style = paragraph.style
            style_element = getattr(style, "element", None)
            if style_element is None:
                return None, None
            # style_element.xml is a string; parse to an element to query
            root = etree.fromstring(style_element.xml.encode("utf-8")) if isinstance(style_element.xml, str) else None
            if root is None:
                return None, None
            ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
            numPr = root.find('.//w:pPr/w:numPr', namespaces=ns)
            if numPr is None:
                return None, None
            numId_elem = numPr.find('w:numId', namespaces=ns)
            ilvl_elem = numPr.find('w:ilvl', namespaces=ns)
            numId = None
            ilvl = None
            if numId_elem is not None:
                val = numId_elem.get(self.XML_KEY)
                numId = self._str_to_int(val, None)
            if ilvl_elem is not None:
                val = ilvl_elem.get(self.XML_KEY)
                ilvl = self._str_to_int(val, None)
            return numId, ilvl
        except Exception:
            return None, None

    def _get_numbering_root(self, docx_obj: DocxDocument):
        """
        Locate and return the numbering part root element (numbering.xml) or None.
        """
        try:
            for rel in docx_obj.part.rels.values():
                reltype = getattr(rel, "reltype", "")
                if isinstance(reltype, str) and reltype.endswith("/numbering"):
                    target_part = getattr(rel, "target_part", None)
                    if target_part is not None:
                        return getattr(target_part, "_element", None)
        except Exception:
            return None
        return None

    def _get_numFmt(
        self, docx_obj: DocxDocument, numId: Optional[int], ilvl: Optional[int]
    ) -> Optional[str]:
        """
        Return w:numFmt value (e.g., 'decimal', 'bullet') for given numId/ilvl.
        If unavailable, return None.
        """
        try:
            if numId is None or ilvl is None:
                return None
            numbering_root = self._get_numbering_root(docx_obj)
            if numbering_root is None:
                return None
            ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
            num_xpath = f".//w:num[@w:numId='{numId}']/w:abstractNumId"
            abstractNumId_el = numbering_root.find(num_xpath, namespaces=ns)
            if abstractNumId_el is None:
                return None
            abstract_id = abstractNumId_el.get(self.XML_KEY)
            if not abstract_id:
                return None
            base_xpath = f".//w:abstractNum[@w:abstractNumId='{abstract_id}']/w:lvl[@w:ilvl='{ilvl}']"
            numFmt_el = numbering_root.find(base_xpath + "/w:numFmt", namespaces=ns)
            if numFmt_el is None:
                return None
            fmt = numFmt_el.get(self.XML_KEY)
            return str(fmt).lower() if fmt is not None else None
        except Exception:
            return None

    def _build_number_label(self, numId: Optional[int], ilvl: Optional[int], docx_obj: DocxDocument) -> Optional[str]:
        """
        Build the numbering label (e.g., "1.3.2") for a given numId/ilvl using numbering.xml.
        Maintains counters per numId across calls.
        """
        if numId is None or ilvl is None:
            return None
        try:
            numbering_root = self._get_numbering_root(docx_obj)
            if numbering_root is None:
                return None
            ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
            # Find abstractNumId via num
            num_xpath = f".//w:num[@w:numId='{numId}']/w:abstractNumId"
            abstractNumId_el = numbering_root.find(num_xpath, namespaces=ns)
            if abstractNumId_el is None:
                return None
            abstract_id = abstractNumId_el.get(self.XML_KEY)
            if abstract_id is None:
                return None
            # Find lvl for ilvl
            lvl_xpath = f".//w:abstractNum[@w:abstractNumId='{abstract_id}']/w:lvl[@w:ilvl='{ilvl}']/w:lvlText"
            lvlText_el = numbering_root.find(lvl_xpath, namespaces=ns)
            if lvlText_el is None:
                return None
            pattern = lvlText_el.get(self.XML_KEY) or ""
            if not hasattr(self, "_num_counters_by_numid"):
                self._num_counters_by_numid = {}
            counters = self._num_counters_by_numid.get(numId)
            if counters is None:
                counters = [0] * 10
                self._num_counters_by_numid[numId] = counters
            # Increment current level and reset deeper levels
            level_idx = max(0, int(ilvl))
            counters[level_idx] += 1
            for j in range(level_idx + 1, len(counters)):
                counters[j] = 0
            # Replace %n placeholders
            def repl(m):
                try:
                    idx = int(m.group(1)) - 1
                    if 0 <= idx < len(counters):
                        return str(counters[idx]) if counters[idx] > 0 else ""
                except Exception:
                    pass
                return ""
            label = re.sub(r"%(\d+)", repl, pattern)
            return label.strip()
        except Exception:
            return None

    def _is_numbered_list(self, paragraph: Paragraph, docx_obj: DocxDocument) -> bool:
        """
        Paragraphì˜ numId/ilvlë¥¼ ë°”íƒ•ìœ¼ë¡œ numbering.xmlì„ ì¡°íšŒí•´
        í•´ë‹¹ ëª©ë¡ì´ ìˆ«ì/ë¬¸ìì‹ì˜ ìˆœì„œí˜•ì¸ì§€(ordered) ì—¬ë¶€ë¥¼ íŒì •í•œë‹¤.

        íŒì • ê¸°ì¤€:
        - w:numFmt ê°€ 'bullet' ë˜ëŠ” 'none' ì´ë©´ False
        - ê·¸ ì™¸ì˜ numFmt ëŠ” True (ëŒ€ë¶€ë¶„ ìˆœì„œí˜•ì„)
        - numFmt ê°€ ì—†ì„ ê²½ìš° w:lvlText ì˜ "%1", "%2" ê°™ì€ í”Œë ˆì´ìŠ¤í™€ë” ì¡´ì¬ì‹œ True
        - ìœ„ë¥¼ ëª¨ë‘ íŒì •í•˜ì§€ ëª»í•˜ë©´ False
        """
        try:
            numid, ilvl = self._get_numId_and_ilvl(paragraph)
            if numid is None or ilvl is None:
                # ìŠ¤íƒ€ì¼ì— ì˜í•œ ë²ˆí˜¸ ì§€ì •ì¼ ìˆ˜ ìˆìœ¼ë¯€ë¡œ ë³´ì¡° ì¡°íšŒ
                s_numid, s_ilvl = self._get_style_numId_and_ilvl(paragraph)
                if numid is None:
                    numid = s_numid
                if ilvl is None:
                    ilvl = s_ilvl

            if numid is None or ilvl is None:
                return False

            numbering_root = self._get_numbering_root(docx_obj)
            if numbering_root is None:
                return False

            ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
            # numId â†’ abstractNumId
            num_xpath = f".//w:num[@w:numId='{numid}']/w:abstractNumId"
            abstractNumId_el = numbering_root.find(num_xpath, namespaces=ns)
            if abstractNumId_el is None:
                return False
            abstract_id = abstractNumId_el.get(self.XML_KEY)
            if not abstract_id:
                return False

            base_xpath = f".//w:abstractNum[@w:abstractNumId='{abstract_id}']/w:lvl[@w:ilvl='{ilvl}']"
            # 1) ëª…ì‹œì  numFmt ìš°ì„ 
            numFmt_el = numbering_root.find(base_xpath + "/w:numFmt", namespaces=ns)
            if numFmt_el is not None:
                fmt = numFmt_el.get(self.XML_KEY)
                if fmt is None:
                    return False
                fmt_str = str(fmt).lower()
                if fmt_str in ("bullet", "none"):
                    return False
                # bullet/none ì´ì™¸ ëŒ€ë¶€ë¶„ì€ ìˆœì„œí˜•ìœ¼ë¡œ ê°„ì£¼
                return True

            # 2) numFmt ì—†ìœ¼ë©´ lvlText íŒ¨í„´ìœ¼ë¡œ ì¶”ì •
            lvlText_el = numbering_root.find(base_xpath + "/w:lvlText", namespaces=ns)
            if lvlText_el is not None:
                pattern = lvlText_el.get(self.XML_KEY) or ""
                # %n í”Œë ˆì´ìŠ¤í™€ë”ê°€ ìˆìœ¼ë©´ ìˆœì„œí˜•
                if re.search(r"%\d+", pattern):
                    return True
                # ëŒ€í‘œì ì¸ ë¶ˆë¦¿ ê¸°í˜¸ê°€ í¬í•¨ë˜ì–´ ìˆìœ¼ë©´ ë¹„ìˆœì„œí˜•ìœ¼ë¡œ ê°„ì£¼
                if any(ch in pattern for ch in ("â€¢", "â—", "â– ", "â€“", "-", "â—‹", "â–ª")):
                    return False

            return False
        except Exception:
            return False

    def _get_heading_and_level(self, style_label: str) -> tuple[str, Optional[int]]:
        parts = self._split_text_and_number(style_label)

        if len(parts) == 2:
            parts.sort()
            label_str: str = ""
            label_level: Optional[int] = 0
            if parts[0].strip().lower() == "heading":
                label_str = "Heading"
                label_level = self._str_to_int(parts[1], None)
            if parts[1].strip().lower() == "heading":
                label_str = "Heading"
                label_level = self._str_to_int(parts[0], None)
            return label_str, label_level

        return style_label, None

    def _get_label_and_level(self, paragraph: Paragraph) -> tuple[str, Optional[int]]:
        if paragraph.style is None:
            return "Normal", None

        label = paragraph.style.style_id
        name = paragraph.style.name
        base_style_label = None
        base_style_name = None
        if base_style := getattr(paragraph.style, "base_style", None):
            base_style_label = base_style.style_id
            base_style_name = base_style.name

        if label is None:
            return "Normal", None

        if ":" in label:
            parts = label.split(":")
            if len(parts) == 2:
                return parts[0], self._str_to_int(parts[1], None)

        if "heading" in label.lower():
            return self._get_heading_and_level(label)
        if name and "heading" in name.lower():
            return self._get_heading_and_level(name)
        if base_style_label and "heading" in base_style_label.lower():
            return self._get_heading_and_level(base_style_label)
        if base_style_name and "heading" in base_style_name.lower():
            return self._get_heading_and_level(base_style_name)

        return label, None

    @classmethod
    def _get_format_from_run(cls, run: Run) -> Optional[Formatting]:
        # The .bold and .italic properties are booleans, but .underline can be an enum
        # like WD_UNDERLINE.THICK (value 6), so we need to convert it to a boolean
        has_bold = run.bold or False
        has_italic = run.italic or False
        # Convert any non-None underline value to True
        has_underline = bool(run.underline is not None and run.underline)

        return Formatting(
            bold=has_bold,
            italic=has_italic,
            underline=has_underline,
        )
        
    # í´ë˜ìŠ¤ ë‚´ë¶€ì— ìœ í‹¸ ì¶”ê°€
    def _resolve_part_by_rid(self, owner_part, rId):
        """
        owner_part.rels ì—ì„œ rIdë¥¼ ì°¾ì•„ target_partë¥¼ ë°˜í™˜.
        owner_part ëŠ” document.xml, headerX.xml, footerY.xml ë“± í˜„ì¬ XMLì˜ ì†Œìœ  íŒŒíŠ¸.
        """
        if not owner_part or not rId:
            return None
        rel = owner_part.rels.get(rId)
        # print(rel,"rel")
        return getattr(rel, "target_part", None) if rel else None
    def _is_owner_header_footer(self, owner_part) -> bool:
        try:
            pn = getattr(owner_part, "partname", None)
            if pn is None:
                return False
            pn_str = str(pn).lower()
            return "/word/header" in pn_str or "/word/footer" in pn_str
        except Exception:
            return False
    def _process_header_footer_refs(self, sectPr_el, docx_obj, doc, owner_part):
        """
        sectPr ì•ˆì˜ w:headerReference / w:footerReference íƒœê·¸ë¥¼ ë§Œë‚˜ëŠ” ìˆœì„œëŒ€ë¡œ ì²˜ë¦¬.
        ê° referenceì˜ r:idë¥¼ í˜„ì¬ owner_part.relsì—ì„œ í•´ì„í•´ header/footer íŒŒíŠ¸ë¥¼ ë¡œë“œ,
        í•´ë‹¹ íŒŒíŠ¸ì˜ ë£¨íŠ¸ ì—˜ë¦¬ë¨¼íŠ¸ë¥¼ _walk_linear ë¡œ ìˆœíšŒí•œë‹¤.
        """
        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
            "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships"}

        # sectPr ìì‹ë“¤ì„ "ë“±ì¥ ìˆœì„œ"ëŒ€ë¡œ ìˆœíšŒí•˜ë©° header/footerReferenceë¥¼ ë§Œë‚˜ëŠ” ì¦‰ì‹œ ì²˜ë¦¬
        for child in sectPr_el.iterchildren():
            lname = etree.QName(child).localname
            # print(lname,"lname")
            if lname in ("headerReference", "footerReference"):
                # print(child,"child")
                rid = child.get("{%s}id" % ns["r"])  # r:id
                # print(rid,"rid")
                target_part = self._resolve_part_by_rid(owner_part, rid)
                # print(target_part,"target_part")
                if not target_part:
                    continue

                root_el = getattr(target_part, "_element", None)  # headerX.xml / footerY.xml ë£¨íŠ¸
                if root_el is None:
                    continue
                # print(root_el,"root_el")
                # â© header/footerì˜ ë‚´ë¶€ë„ "ê·¸ ìë¦¬ì—ì„œ" ìˆœì°¨ íŒŒì‹±
                self._walk_linear(
                    body=root_el,
                    docx_obj=docx_obj,
                    doc=doc,
                    owner_part=target_part,
            )    
    def _walk_linear(
        self,
        body: BaseOxmlElement,
        docx_obj: DocxDocument,
        doc: DoclingDocument,
        owner_part=None,  # â† ì¶”ê°€: í˜„ì¬ bodyë¥¼ ì†Œìœ í•œ OPC íŒŒíŠ¸
    ) -> DoclingDocument:
        if owner_part is None:
            owner_part = docx_obj.part
        # Header/Footer ë£¨íŠ¸ì¼ ê²½ìš° ì„¹ì…˜ ê·¸ë£¹ ì»¨í…ìŠ¤íŠ¸ë¥¼ ì—´ì–´ì¤€ë‹¤
        header_footer_ctx_opened = False
        original_parent_for_hf = None
        try:
            root_local = etree.QName(body).localname
        except Exception:
            root_local = None
        if root_local in ("hdr", "ftr"):
            level = self._get_level()
            group_name = "header" if root_local == "hdr" else "footer"
            hf_group = doc.add_group(
                label=GroupLabel.SECTION,
                parent=self.parents.get(level - 1),
                name=group_name,
            )
            original_parent_for_hf = self.parents.get(level)
            # ë˜í•œ level-1 ë¶€ëª¨ë„ ì„ì‹œë¡œ í—¤ë”/í‘¸í„° ê·¸ë£¹ìœ¼ë¡œ ë®ì–´ì¨ì„œ ë‚´ë¶€ ì¶”ê°€ë¬¼ì´ ê·¸ë£¹ì— ê·€ì†ë˜ë„ë¡ í•¨
            original_parent_for_hf_m1 = self.parents.get(level - 1)
            self.parents[level - 1] = hf_group
            self.parents[level] = hf_group
            header_footer_ctx_opened = True
            try:
                print(f"[HF] enter {group_name}")
            except Exception:
                pass
        for element in body:

            # Check for Inline Images (blip elements)
            namespaces = {
                "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
                "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
                "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
                "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
                "mc": "http://schemas.openxmlformats.org/markup-compatibility/2006",
                "v": "urn:schemas-microsoft-com:vml",
                "wps": "http://schemas.microsoft.com/office/word/2010/wordprocessingShape",
                "w10": "urn:schemas-microsoft-com:office:word",
                "a14": "http://schemas.microsoft.com/office/drawing/2010/main",
            }
            xpath_expr = etree.XPath(".//a:blip", namespaces=namespaces)
            drawing_blip = xpath_expr(element)
            # Skip the fallback inside mc:AlternateContent
            tag = etree.QName(element).localname
            if tag == "AlternateContent":
                # find the mc:Choice branch and process only that
                choice = element.find("mc:Choice", namespaces=namespaces)
                if choice is not None:
                    # inline its children into our loop
                    for child in choice:
                        doc = self._walk_linear(child, docx_obj, doc, owner_part)
                # skip the rest (Fallback)
                continue
            tag_name = etree.QName(element).localname
            if header_footer_ctx_opened:
                try:
                    print(f"[HF] tag={tag_name}")
                except Exception:
                    pass
            # 1) body ì§ì† sectPr (ë“œë¬¼ì§€ë§Œ ì¡´ì¬)
            sectprs = []
            if tag_name == "sectPr":
                sectprs = [element]

            # 2) ì¼ë°˜ ì¼€ì´ìŠ¤: ë¬¸ë‹¨ pPr ì•ˆì˜ sectPr
            elif tag_name == "p":
                sectprs = element.findall("./w:pPr/w:sectPr", namespaces=namespaces)

            # 3) ë§Œë‚œ ìˆœì„œëŒ€ë¡œ ì „ë¶€ ì²˜ë¦¬ (ì¤‘ë³µ ë°©ì§€)
            for sectPr_el in sectprs:
                sid = id(sectPr_el)
                if sid in self._seen_sectpr_ids:
                    continue
                self._seen_sectpr_ids.add(sid)

                self._process_header_footer_refs(
                    sectPr_el=sectPr_el,
                    docx_obj=docx_obj,
                    doc=doc,
                    owner_part=owner_part,
                )
            # Check for shape content (including textboxes and other shapes)
            # Only process if the element hasn't been processed before
            element_id = id(element)
            if element_id not in self.processed_textbox_elements:
                # Modern Word textboxes
                txbx_xpath = etree.XPath(
                    ".//w:txbxContent|.//v:textbox//w:p", namespaces=namespaces
                )
                textbox_elements = txbx_xpath(element)

                # No modern textboxes found, check for alternate/legacy textbox formats
                if not textbox_elements and tag_name in ["drawing", "pict"]:
                    # Additional checks for textboxes in DrawingML and VML formats
                    alt_txbx_xpath = etree.XPath(
                        ".//wps:txbx//w:p|.//w10:wrap//w:p|.//a:p//a:t",
                        namespaces=namespaces,
                    )
                    textbox_elements = alt_txbx_xpath(element)

                    # Check for shape text that's not in a standard textbox
                    if not textbox_elements:
                        shape_text_xpath = etree.XPath(
                            ".//a:bodyPr/ancestor::*//a:t|.//a:txBody//a:t",
                            namespaces=namespaces,
                        )
                        shape_text_elements = shape_text_xpath(element)
                        if shape_text_elements:
                            # Create custom text elements from shape text
                            text_content = " ".join(
                                [t.text for t in shape_text_elements if t.text]
                            )
                            if text_content.strip():
                                # Create a paragraph-like element to process with standard handler
                                level = self._get_level()
                                shape_group = doc.add_group(
                                    label=GroupLabel.SECTION,
                                    parent=self.parents[level - 1],
                                    name="shape-text",
                                )
                                doc.add_text(
                                    label=DocItemLabel.PARAGRAPH,
                                    parent=shape_group,
                                    text=text_content,
                                    prov=ProvenanceItem(
                                    page_no=1,
                                    bbox=BoundingBox(l=0, t=0, r=1, b=1),
                                    charspan=(0, 0)
                                )
                                )

                if textbox_elements:
                    # Mark the parent element as processed
                    self.processed_textbox_elements.append(element_id)
                    # Also mark all found textbox elements as processed
                    for tb_element in textbox_elements:
                        self.processed_textbox_elements.append(id(tb_element))

                    self._handle_textbox_content(textbox_elements, docx_obj, doc)

            # Check for shape content (similar to hwpx_backend's _process_rect)
            if tag_name in ["drawing", "pict"] and element_id not in self.processed_textbox_elements:
                self._handle_shape_content(element, docx_obj, doc, owner_part=owner_part)

            # Check for Tables - Use enhanced table processing
            if element.tag.endswith("tbl"):
                try:
                    if header_footer_ctx_opened:
                        try:
                            print("[HF] handle table")
                        except Exception:
                            pass
                    # header/footer ë‚´ë¶€ í…Œì´ë¸” ì´ë¯¸ì§€ë¥¼ ìœ„í•´ owner_part ì „ë‹¬
                    self._handle_tables_enhanced(element, docx_obj, doc, owner_part=owner_part)
                except Exception as e:
                    _log.debug(f"[MSWORD] _handle_tables_enhanced failed: {e}")

            elif drawing_blip:
                # ğŸ” ì†Œìœ  íŒŒíŠ¸ë¥¼ ë„˜ê²¨ì„œ rId í•´ì„ì´ ì˜¬ë°”ë¥´ê²Œ header/footerì—ì„œë„ ì‘ë™
                if header_footer_ctx_opened:
                    try:
                        print("[HF] handle image (blip)", drawing_blip)
                    except Exception:
                        pass
                self._handle_pictures(owner_part, docx_obj, drawing_blip, doc)
                # ì´ë¯¸ì§€ ë’¤ í…ìŠ¤íŠ¸ ì²˜ë¦¬
                if (tag_name in ["p"]
                    and (element.find(".//w:t", namespaces=namespaces) is not None or element.find(".//w:instrText", namespaces=namespaces) is not None)):
                    if header_footer_ctx_opened:
                        try:
                            texts = [t.text for t in element.findall(".//w:t", namespaces=namespaces) if t.text]
                            instrs = [t.text for t in element.findall(".//w:instrText", namespaces=namespaces) if t.text]
                            fldchars = [fc.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fldCharType") for fc in element.findall(".//w:fldChar", namespaces=namespaces)]
                            parts = []
                            if texts:
                                parts.append(" ".join(texts))
                            if instrs:
                                parts.append(f"[instr] {' '.join(instrs)}")
                            if fldchars:
                                parts.append(f"[fld] {','.join([c for c in fldchars if c])}")
                            joined = " | ".join(parts)
                            preview = (joined[:80] + ("â€¦" if len(joined) > 80 else "")) if joined else ""
                            # print(f"[HF] handle paragraph after image text='{preview}'")
                        except Exception:
                            pass
                    self._handle_text_elements(element, docx_obj, doc)    
                              
            # Check for the sdt containers, like table of contents
            elif tag_name in ["sdt"]:
                sdt_content = element.find(".//w:sdtContent", namespaces=namespaces)
                if sdt_content is not None:
                    if header_footer_ctx_opened:
                        try:
                            print("[HF] handle sdt content")
                        except Exception:
                            pass
                    paragraphs = sdt_content.findall(".//w:p", namespaces=namespaces)
                    for p in paragraphs:
                        self._handle_text_elements(p, docx_obj, doc)
            # Check for Text
            elif tag_name in ["p"]:
                # "tcPr", "sectPr"
                if header_footer_ctx_opened:
                    try:
                        texts = [t.text for t in element.findall(".//w:t", namespaces=namespaces) if t.text]
                        instrs = [t.text for t in element.findall(".//w:instrText", namespaces=namespaces) if t.text]
                        fldchars = [fc.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fldCharType") for fc in element.findall(".//w:fldChar", namespaces=namespaces)]
                        parts = []
                        if texts:
                            parts.append(" ".join(texts))
                        if instrs:
                            parts.append(f"[instr] {' '.join(instrs)}")
                        if fldchars:
                            parts.append(f"[fld] {','.join([c for c in fldchars if c])}")
                        joined = " | ".join(parts)
                        preview = (joined[:80] + ("â€¦" if len(joined) > 80 else "")) if joined else ""
                        # print(f"[HF] handle paragraph text='{preview}'")
                    except Exception:
                        pass
                self._handle_text_elements(element, docx_obj, doc)
                
        # Header/Footer ì»¨í…ìŠ¤íŠ¸ ë‹«ê¸°
        if header_footer_ctx_opened:
            level = self._get_level()
            self.parents[level] = original_parent_for_hf
            # level-1 ë¶€ëª¨ ë³µì›
            try:
                self.parents[level - 1] = original_parent_for_hf_m1
            except Exception:
                pass
            try:
                print(f"[HF] exit {group_name}")
            except Exception:
                pass
        return doc 

    def _extract_image_from_drawing(
        self, drawing_el: BaseOxmlElement, docx_obj: DocxDocument
    ) -> Optional[ImageRef]:
        """
        <w:drawing> í˜¹ì€ VML <v:imagedata> ê°™ì€ ìš”ì†Œì—ì„œ
        Word ê´€ê³„(rId)ë¥¼ ì°¾ì•„ ì´ë¯¸ì§€ë¥¼ ì¶”ì¶œí•©ë‹ˆë‹¤.
        """
        # Word ML namespace
        ns = {
            "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
            "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
            "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
            "pic": "http://schemas.openxmlformats.org/drawingml/2006/picture",
        }

        # 1) <a:blip> ì°¾ê¸°
        blip = drawing_el.find(".//a:blip", namespaces=ns)
        if blip is None:
            return None

        # 2) ê´€ê³„ ID ì¶”ì¶œ
        embed_rId = blip.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
        if not embed_rId or embed_rId not in docx_obj.part.rels:
            return None

        # 3) ì´ë¯¸ì§€ ë°”ì´ë„ˆë¦¬ ê°€ì ¸ì˜¤ê¸°
        image_part = docx_obj.part.rels[embed_rId].target_part
        blob = image_part.blob
        try:
            pil_img = Image.open(BytesIO(blob))
        except UnidentifiedImageError:
            return None

        # 4) ImageRef ìƒì„±
        return ImageRef.from_pil(image=pil_img, dpi=72)
    
    def _handle_tables_enhanced(self, element: etree._Element, docx_obj: DocxDocument, doc: DoclingDocument, owner_part=None) -> None:
        # ì´ element ê°€ mc:Fallback ê³„ì¸µ ì•ˆì´ë¼ë©´ ìŠ¤í‚µ
        if owner_part is None:
            owner_part = docx_obj.part
        hf_only = self._is_owner_header_footer(owner_part)
        if element.getparent() is not None and etree.QName(element.getparent()).localname == "Fallback":
            if hf_only:
                try:
                    print("[HF][tbl] skip: inside Fallback")
                except Exception:
                    pass
            return

        # ë³´ìˆ˜ì ì¸ ì¤‘ë³µ ì œê±°: mc:AlternateContent ë‚´ë¶€ì˜ ëª…í™•í•œ ì¤‘ë³µë§Œ ì œê±°
        parent = element.getparent()
        is_in_alternate_content = False
        
        # Check if this table is inside mc:AlternateContent structure
        ancestor = parent
        while ancestor is not None:
            if etree.QName(ancestor).localname == "AlternateContent":
                is_in_alternate_content = True
                break
            ancestor = ancestor.getparent()
        
        # Only apply content hash duplicate detection for mc:AlternateContent tables
        if is_in_alternate_content:
            try:
                table = Table(element, docx_obj)
                table_content_hash = self._get_table_content_hash(table)
                
                if not hasattr(self, '_processed_table_contents'):
                    self._processed_table_contents = set()
                
                if table_content_hash in self._processed_table_contents:
                    if hf_only:
                        try:
                            print("[HF][tbl] skip: duplicate content in AlternateContent")
                        except Exception:
                            pass
                    return
                
                self._processed_table_contents.add(table_content_hash)
                
            except Exception as e:
                _log.debug(f"[MSWORD] table content hash failed: {e}")

        eid = id(element)
        #--- í…Œì´ë¸” ë‚´ë¶€ì˜ í…ìŠ¤íŠ¸ ì „ì²´ í™•ì¸
        # python-docx Table ê°ì²´ë¡œ ë³€í™˜í•´ì„œ ê° ì…€ì„ ìˆœíšŒ
        table = Table(element, docx_obj)

        # ì¤‘ë³µ ë°©ì§€ë¥¼ ìœ„í•œ ID ì²´í¬
        if eid in self.processed_table_elements:
            return
        self.processed_table_elements.add(eid)
                        
        # 1) ê¸°ë³¸ í…Œì´ë¸” í¬ê¸°
        table = Table(element, docx_obj)
        num_rows = len(table.rows)
        num_cols = len(table.columns)
        if hf_only:
            try:
                print(f"[HF][tbl] enter rows={num_rows} cols={num_cols}")
            except Exception:
                pass
        ns = element.nsmap

        # 2) table-level detection: ì¤‘ì²© tbl / ê·¸ë¦¼ì´ ìˆëŠ”ì§€
        #    - [0]ë²ˆì§¸ëŠ” ìê¸° ìì‹ (<w:tbl>)ì´ ì¡íˆê¸° ë•Œë¬¸ì— [1:]ë¡œ ì‹¤ì œ ì¤‘ì²© í…Œì´ë¸”ë§Œ
        nested_tbls_global = element.findall('.//w:tbl', namespaces=ns)[1:]
        pics_global = (
            element.findall('.//w:drawing', namespaces=ns) +
            element.findall('.//v:imagedata', namespaces=ns)
        )
        table_has_nested = bool(nested_tbls_global)
        table_has_pics   = bool(pics_global)
        if num_rows == 1 and num_cols == 1:
            cell_element = table.rows[0].cells[0]
            # In case we have a table of only 1 cell, we consider it furniture
            # And proceed processing the content of the cell as though it's in the document body
            if hf_only:
                try:
                    print("[HF][tbl] 1x1 furniture: inline its content")
                except Exception:
                    pass
            self._walk_linear(cell_element._element, docx_obj, doc)
            return
        # 2) ìˆœìˆ˜ TableDataë¥¼ ìŒ“ì„ ê°ì²´
        data = TableData(num_rows=num_rows, num_cols=num_cols)

        # 3) ì¤‘ì²© êµ¬ì¡° ë²„í¼: (r, c) â†’ list of (typ, payload)
        cell_buffer = defaultdict(list)
            
        def get_docx_image_bytes_from_owner(drawing_blip: List[etree._Element]) -> Optional[bytes]:
            rId = drawing_blip[0].get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
            if not rId or rId not in owner_part.rels:
                return None
            return owner_part.rels[rId].target_part.blob
        
        # 4) ê° ì…€ ìˆœíšŒ
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                
                # ì§ì† ìì‹ë§Œ
                children = list(cell._element.getchildren())
                # ê°™ì€ depthì—ì„œ ì¤‘ì²© tbl / ê·¸ë¦¼ ê²€ì‚¬
                # nested_tbls = [ch for ch in children if etree.QName(ch).localname == "tbl"]
                
                # pics        = cell._element.findall(".//w:drawing", namespaces=ns) \
                #             + cell._element.findall(".//v:imagedata", namespaces=ns) 

                if table_has_nested or table_has_pics:
                    # ì¤‘ì²© ìˆëŠ” ì…€: ìì‹ ìˆœì„œëŒ€ë¡œ buffer ì €ì¥
                    for ch in children:
                        tag = etree.QName(ch).localname
                        if tag == "tbl":
                            cell_buffer[(r_idx, c_idx)].append(("table", ch))
                            if hf_only:
                                try:
                                    print(f"[HF][tbl] cell({r_idx},{c_idx}) -> nested table")
                                except Exception:
                                    pass
                            continue
                            
                        elif tag == "p":
                            # -- 1) í…ìŠ¤íŠ¸ ìˆ˜ì§‘
                            texts = [
                                t.text.strip()
                                for t in ch.findall(".//w:t", namespaces=ns)
                                if t.text and t.text.strip()
                            ]
                            if texts:
                                if not self._is_duplicate_content(" ".join(texts)):
                                    cell_buffer[(r_idx, c_idx)].append(("text", " ".join(texts))) 
                                    if hf_only:
                                        try:
                                            preview = (" ".join(texts))[:80]
                                            # print(f"[HF][tbl] cell({r_idx},{c_idx}) -> text '{preview}{'â€¦' if len(' '.join(texts))>80 else ''}'")
                                        except Exception:
                                            pass
                                continue

                            # -- 2) drawing ìˆ˜ì§‘
                            drawings = ch.findall(".//w:drawing", namespaces=ns)
                            if drawings:
                                blob = get_docx_image_bytes_from_owner(drawings)
                                if blob is None:
                                    cell_buffer[(r_idx, c_idx)].append(("picture", None))
                                    if hf_only:
                                        try:
                                            print(f"[HF][tbl] cell({r_idx},{c_idx}) -> picture (no blob)")
                                        except Exception:
                                            pass
                                else:
                                    try:
                                        pil_img = Image.open(BytesIO(blob))
                                        img_ref = ImageRef.from_pil(image=pil_img, dpi=72)
                                        cell_buffer[(r_idx, c_idx)].append(("picture", img_ref))
                                        if hf_only:
                                            try:
                                                print(f"[HF][tbl] cell({r_idx},{c_idx}) -> picture OK")
                                            except Exception:
                                                pass
                                    except UnidentifiedImageError:
                                        # ì‹¤íŒ¨í•´ë„ ìë¦¬ í‘œì‹œ
                                        cell_buffer[(r_idx, c_idx)].append(("picture", None)) 
                                        if hf_only:
                                            try:
                                                print(f"[HF][tbl] cell({r_idx},{c_idx}) -> picture unreadable")
                                            except Exception:
                                                pass
                            continue                                


                    # ë§Œì•½ ë²„í¼ì— ë­”ê°€ ë‹´ê²¼ë‹¤ë©´, TableDataì— ì¶”ê°€í•˜ì§€ ì•Šê³  continue
                    if (r_idx, c_idx) in cell_buffer:
                        if hf_only:
                            try:
                                print(f"[HF][tbl] cell({r_idx},{c_idx}) buffered-only (skip TableData)")
                            except Exception:
                                pass
                        continue
                    
                # 5) ì¼ë°˜ ì…€: TableData ì— ì¶”ê°€
                cell_text = self._extract_cell_text_with_sdt(cell).strip()
                data.table_cells.append(
                    TableCell(
                        text=cell_text,
                        row_span=1,
                        col_span=1,
                        start_row_offset_idx=r_idx,
                        end_row_offset_idx=r_idx + 1,
                        start_col_offset_idx=c_idx,
                        end_col_offset_idx=c_idx + 1,
                        column_header=(r_idx == 0),
                        row_header=False,
                    )
                )
                if hf_only:
                    try:
                        preview = cell_text[:80]
                        print(f"[HF][tbl] cell({r_idx},{c_idx}) -> TableData '{preview}{'â€¦' if len(cell_text)>80 else ''}'")
                    except Exception:
                        pass

        # 6) ë²„í¼ ì¶œë ¥
        parent = self.parents[self._get_level() - 1]
        if parent is None:
            # Try current level
            parent = self.parents.get(self._get_level())
            # Try to find nearest non-None parent (prefer HF group)
            if parent is None:
                for k in range(self._get_level(), -2, -1):
                    if self.parents.get(k) is not None:
                        parent = self.parents[k]
                        break
            # Debug
            if self._is_owner_header_footer(owner_part):
                try:
                    pname = type(parent).__name__ if parent is not None else "None"
                    print(f"[HF][tbl] parent fallback -> {pname}")
                except Exception:
                    pass
        for (r, c), items in sorted(cell_buffer.items(), key=lambda x: (x[0][0], x[0][1])):
            for typ, payload in items:
                prov = ProvenanceItem(
                    page_no=1,
                    bbox=BoundingBox(l=0, t=r, r=1, b=r+1),
                    charspan=(0, len(payload) if typ == "text" else 0)
                )
                if typ == "text":
                    doc.add_text(label=DocItemLabel.PARAGRAPH, text=payload, parent=parent, prov=prov)
                    if hf_only:
                        try:
                            pvw = payload[:80]
                            print(f"[HF][tbl] emit text from buffer ({r},{c}) '{pvw}{'â€¦' if len(payload)>80 else ''}'")
                        except Exception:
                            pass
                    continue
                elif typ == "picture":
                    if payload is None:
                        if hf_only:
                            try:
                                print(f"[HF][tbl] skip picture (no payload) ({r},{c})")
                            except Exception:
                                pass
                        continue  # ìë¦¬ í‘œì‹œê°€ ì—†ìœ¼ë©´ ìŠ¤í‚µ
                    else:
                        doc.add_picture(parent=parent, image=payload, caption=None, prov=prov)
                        if hf_only:
                            try:
                                print(f"[HF][tbl] emit picture ({r},{c})")
                            except Exception:
                                pass
                        continue
                elif typ == "table":
                    # ì¤‘ì²© í…Œì´ë¸”(ê°€ì¥ ì•ˆìª½)ë§Œ ì‹¤ì œ TableDataë¡œ ì¬ê·€ ì²˜ë¦¬
                    if hf_only:
                        try:
                            print(f"[HF][tbl] recurse nested table ({r},{c})")
                        except Exception:
                            pass
                    self._handle_tables_enhanced(payload, docx_obj, doc)
                    continue

        # 7) TableData í˜•íƒœë¡œ ì¶œë ¥ (ì¤‘ì²© ì—†ëŠ” ê°€ì¥ ë°”ê¹¥ìª½ë§Œ)
        if data.table_cells:
            # parent ë³´ì •: ì—¬ì „íˆ None ì´ë©´ ìµœìƒìœ„ì— ë¶™ì—¬ ì¤‘ë‹¨ì„ ë°©ì§€
            if parent is None:
                if self._is_owner_header_footer(owner_part):
                    try:
                        print("[HF][tbl] warn: parent None, attaching table at document root")
                    except Exception:
                        pass
                parent = self.parents.get(0)
            doc.add_table(data=data, parent=parent, prov=ProvenanceItem(
                page_no=1,
                bbox=BoundingBox(l=0, t=0, r=1, b=num_rows),
                charspan=(0, 0)
            ))
            if hf_only:
                try:
                    print(f"[HF][tbl] emit TableData rows={num_rows} cols={num_cols} cells={len(data.table_cells)}")
                except Exception:
                    pass
        else:
            if hf_only:
                try:
                    print("[HF][tbl] no TableData cells emitted")
                except Exception:
                    pass

    def _should_fallback_table_to_text(self, table: Table, docx_obj: DocxDocument) -> bool:
        """
        Determine if the table should be processed as text instead of a table.
        Returns True if table has complex nested structures that are better processed as text.
        """
        # Check for 1x1 tables
        if len(table.rows) == 1 and len(table.columns) == 1:
            return True
            
        complex_structure_count = 0
        
        for row in table.rows:
            for cell in row.cells:
                # Check for nested tables (currently disabled)
                nested_tables = self._find_nested_tables_in_cell(cell)
                if nested_tables:
                    complex_structure_count += len(nested_tables)
                
                # Check for images (pic elements)
                namespaces = {
                    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
                    "pic": "http://schemas.openxmlformats.org/drawingml/2006/picture",
                }
                
                # Look for picture elements
                xpath_expr = etree.XPath(".//pic:pic", namespaces=namespaces)
                pics = xpath_expr(cell._element)
                if pics:
                    complex_structure_count += len(pics)
                    
                # Look for drawing elements which often contain images
                xpath_expr = etree.XPath(".//w:drawing", namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"})
                drawings = xpath_expr(cell._element)
                if drawings:
                    complex_structure_count += len(drawings)
        
        # If we find multiple complex structures, fallback to text
        return complex_structure_count >= 2

    def _pop_table_to_text(self, table: Table, doc: DoclingDocument) -> None:
        """
        Pop current table and process its content as regular text.
        This is the core pop mechanism requested by the user.
        """
        level = self._get_level()
        parent = self.parents[level - 1] if level > 0 else None
        
        all_text_parts = []
        
        for row in table.rows:
            row_text_parts = []
            for cell in row.cells:
                cell_text = self._extract_cell_text_with_sdt(cell).strip()
                if cell_text:
                    row_text_parts.append(cell_text)
            
            if row_text_parts:
                # Join cell contents with spaces for each row
                row_text = " ".join(row_text_parts)
                all_text_parts.append(row_text)
        
        # Combine all row texts and add as regular text
        if all_text_parts:
            final_text = "\n".join(all_text_parts)
            
            # Check for duplicate content before adding
            if not self._is_duplicate_content(final_text):
                doc.add_text(
                    label=DocItemLabel.TEXT,
                    text=final_text,
                    parent=parent,
                    prov=ProvenanceItem(
                        page_no=1,
                        bbox=BoundingBox(l=0, t=0, r=1, b=1),
                        charspan=(0, len(final_text))
                    )
                )
    

    def _process_table_as_text(self, table: Table, doc: DoclingDocument) -> None:
        """
        Process table content as regular text instead of table structure.
        """
        level = self._get_level()
        parent = self.parents[level - 1] if level > 0 else None
        
        all_text_parts = []
        
        for row in table.rows:
            row_text_parts = []
            for cell in row.cells:
                cell_text = self._extract_cell_text_with_sdt(cell).strip()
                if cell_text:
                    row_text_parts.append(cell_text)
            
            if row_text_parts:
                # Join cell contents with spaces for each row
                row_text = " ".join(row_text_parts)
                all_text_parts.append(row_text)
        
        # Combine all row texts
        if all_text_parts:
            full_text = "\n".join(all_text_parts)
            doc.add_text(
                label=DocItemLabel.TEXT,
                text=full_text,
                parent=parent,
                prov=ProvenanceItem(
                    page_no=1,
                    bbox=BoundingBox(l=0, t=0, r=1, b=1),
                    charspan=(0, len(full_text))
                )
            )

    def _find_nested_tables_in_cell(self, cell: _Cell) -> List[BaseOxmlElement]:
        """Find nested tables within a cell. Temporarily disabled to fix pop issue."""
        # TEMPORARY: Disable nested table detection to fix the comprehensive income statement issue
        # Real nested tables are very rare in practice, and the current detection was causing
        # adjacent tables to be incorrectly identified as nested tables
        return []

    def _find_images_in_cell(self, cell: _Cell, docx_obj: DocxDocument) -> List[Any]:
        """Find images within a cell."""
        images = []
        namespaces = {
            "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
            "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
        }
        
        # Look for blip elements (images)
        xpath_expr = etree.XPath(".//a:blip", namespaces=namespaces)
        blips = xpath_expr(cell._element)
        images.extend(blips)
        
        return images

    def _get_table_content_hash(self, table: Table) -> str:
        """
        Generate a hash of table content to detect duplicate tables.
        """
        import hashlib
        
        content_parts = []
        for row in table.rows:
            row_parts = []
            for cell in row.cells:
                cell_text = self._extract_cell_text_with_sdt(cell).strip()
                row_parts.append(cell_text)
            content_parts.append("|".join(row_parts))
        
        table_content = "\n".join(content_parts)
        return hashlib.md5(table_content.encode('utf-8')).hexdigest()

    def _get_text_content_hash(self, text: str) -> str:
        """
        Generate a hash of text content to detect duplicate content.
        """
        import hashlib
        
        # Normalize text: remove extra whitespace and convert to lowercase
        normalized_text = re.sub(r'\s+', ' ', text.strip().lower())
        return hashlib.md5(normalized_text.encode('utf-8')).hexdigest()

    def _is_duplicate_content(self, text: str) -> bool:
        """
        Check if text content is duplicate based on hash.
        """
        if not text or len(text.strip()) < 5:  # Skip very short texts
            return False
            
        text_hash = self._get_text_content_hash(text)
        
        if not hasattr(self, "_processed_text_contents"):
            self._processed_text_contents = set()
            
        if text_hash in self._processed_text_contents:
            return True
            
        self._processed_text_contents.add(text_hash)
        return False

    def _extract_cell_text_with_sdt(self, cell: _Cell) -> str:
        """
        Extract text from cell including content inside SDT (Structured Document Tags).
        This is needed for cells that have form fields or content controls.
        """
        namespaces = {
            "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        }
        
        # First try the normal cell text
        cell_text = cell.text
        
        # If cell text is empty or very short, look for SDT content
        if not cell_text or len(cell_text.strip()) < 3:
            # Look for all text elements including those inside SDT
            xpath_expr = etree.XPath(".//w:t", namespaces=namespaces)
            text_elements = xpath_expr(cell._element)
            
            all_texts = []
            for t_elem in text_elements:
                if t_elem.text:
                    all_texts.append(t_elem.text)
            
            if all_texts:
                cell_text = "".join(all_texts)
        
        return cell_text if cell_text else "" 

    def _get_paragraph_elements(self, paragraph: Paragraph):
        """
        Extract paragraph elements along with their formatting and hyperlink
        """

        # for now retain empty paragraphs for backwards compatibility:
        if paragraph.text.strip() == "":
            return [("", None, None)]

        paragraph_elements: list[
            tuple[str, Optional[Formatting], Optional[Union[AnyUrl, Path]]]
        ] = []
        group_text = ""
        previous_format = None

        # Iterate over the runs of the paragraph and group them by format
        for c in paragraph.iter_inner_content():
            if isinstance(c, Hyperlink):
                text = c.text
                hyperlink = Path(c.address)
                format = self._get_format_from_run(c.runs[0])
            elif isinstance(c, Run):
                text = c.text
                hyperlink = None
                format = self._get_format_from_run(c)
            else:
                continue

            if (len(text.strip()) and format != previous_format) or (
                hyperlink is not None
            ):
                # If the style changes for a non empty text, add the previous group
                if len(group_text.strip()) > 0:
                    paragraph_elements.append(
                        (group_text.strip(), previous_format, None)
                    )
                group_text = ""

                # If there is a hyperlink, add it immediately
                if hyperlink is not None:
                    paragraph_elements.append((text.strip(), format, hyperlink))
                    text = ""
                else:
                    previous_format = format

            group_text += text

        # Format the last group
        if len(group_text.strip()) > 0:
            paragraph_elements.append((group_text.strip(), format, None))

        return paragraph_elements

    def _get_paragraph_position(self, paragraph_element):
        """Extract vertical position information from paragraph element."""
        # First try to directly get the index from w:p element that has an order-related attribute
        if (
            hasattr(paragraph_element, "getparent")
            and paragraph_element.getparent() is not None
        ):
            parent = paragraph_element.getparent()
            # Get all paragraph siblings
            paragraphs = [
                p for p in parent.getchildren() if etree.QName(p).localname == "p"
            ]
            # Find index of current paragraph within its siblings
            try:
                paragraph_index = paragraphs.index(paragraph_element)
                return paragraph_index  # Use index as position for consistent ordering
            except ValueError:
                pass

        # Look for position hints in element attributes and ancestor elements
        for elem in (*[paragraph_element], *paragraph_element.iterancestors()):
            # Check for direct position attributes
            for attr_name in ["y", "top", "positionY", "y-position", "position"]:
                value = elem.get(attr_name)
                if value:
                    try:
                        # Remove any non-numeric characters (like 'pt', 'px', etc.)
                        clean_value = re.sub(r"[^0-9.]", "", value)
                        if clean_value:
                            return float(clean_value)
                    except (ValueError, TypeError):
                        pass

            # Check for position in transform attribute
            transform = elem.get("transform")
            if transform:
                # Extract translation component from transform matrix
                match = re.search(r"translate\([^,]+,\s*([0-9.]+)", transform)
                if match:
                    try:
                        return float(match.group(1))
                    except ValueError:
                        pass

            # Check for anchors or relative position indicators in Word format
            # 'dist' attributes can indicate relative positioning
            for attr_name in ["distT", "distB", "anchor", "relativeFrom"]:
                if elem.get(attr_name) is not None:
                    return elem.sourceline  # Use the XML source line number as fallback

        # For VML shapes, look for specific attributes
        for ns_uri in paragraph_element.nsmap.values():
            if "vml" in ns_uri:
                # Try to extract position from style attribute
                style = paragraph_element.get("style")
                if style:
                    match = re.search(r"top:([0-9.]+)pt", style)
                    if match:
                        try:
                            return float(match.group(1))
                        except ValueError:
                            pass

        # If no better position indicator found, use XML source line number as proxy for order
        return (
            paragraph_element.sourceline
            if hasattr(paragraph_element, "sourceline")
            else None
        )

    def _collect_textbox_paragraphs(self, textbox_elements):
        """Collect and organize paragraphs from textbox elements."""
        processed_paragraphs = []
        container_paragraphs = {}

        for element in textbox_elements:
            element_id = id(element)
            # Skip if we've already processed this exact element
            if element_id in processed_paragraphs:
                continue

            tag_name = etree.QName(element).localname
            processed_paragraphs.append(element_id)

            # Handle paragraphs directly found (VML textboxes)
            if tag_name == "p":
                # Find the containing textbox or shape element
                container_id = None
                for ancestor in element.iterancestors():
                    if any(ns in ancestor.tag for ns in ["textbox", "shape", "txbx"]):
                        container_id = id(ancestor)
                        break

                if container_id not in container_paragraphs:
                    container_paragraphs[container_id] = []
                container_paragraphs[container_id].append(
                    (element, self._get_paragraph_position(element))
                )

            # Handle txbxContent elements (Word DrawingML textboxes)
            elif tag_name == "txbxContent":
                paragraphs = element.findall(".//w:p", namespaces=element.nsmap)
                container_id = id(element)
                if container_id not in container_paragraphs:
                    container_paragraphs[container_id] = []

                for p in paragraphs:
                    p_id = id(p)
                    if p_id not in processed_paragraphs:
                        processed_paragraphs.append(p_id)
                        container_paragraphs[container_id].append(
                            (p, self._get_paragraph_position(p))
                        )
            else:
                # Try to extract any paragraphs from unknown elements
                paragraphs = element.findall(".//w:p", namespaces=element.nsmap)
                container_id = id(element)
                if container_id not in container_paragraphs:
                    container_paragraphs[container_id] = []

                for p in paragraphs:
                    p_id = id(p)
                    if p_id not in processed_paragraphs:
                        processed_paragraphs.append(p_id)
                        container_paragraphs[container_id].append(
                            (p, self._get_paragraph_position(p))
                        )

        return container_paragraphs

    def _handle_textbox_content(
        self,
        textbox_elements: list,
        docx_obj: DocxDocument,
        doc: DoclingDocument,
    ) -> None:
        level = self._get_level()
        textbox_group = doc.add_group(
            label=GroupLabel.SECTION,
            parent=self.parents[level - 1],
            name="textbox",
        )
        original_parent = self.parents[level]
        self.parents[level] = textbox_group

        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

        # 1) textbox_elements ì¤‘ ì‹¤ì œ txbxContent ë…¸ë“œë§Œ ë½‘ê¸°
        txbx_contents = [
            el for el in textbox_elements
            if etree.QName(el).localname == "txbxContent"
        ]

        # 2) ê° txbxContentì˜ ìì‹ ë…¸ë“œë¥¼ ìˆœì„œëŒ€ë¡œ ìˆœíšŒ
        for content in txbx_contents:
            for child in content.getchildren():
                local = etree.QName(child).localname

                if local == "p":
                    # ë‹¨ë½ì´ë©´ ë°”ë¡œ í…ìŠ¤íŠ¸ ì²˜ë¦¬
                    self._handle_text_elements(
                        child, docx_obj, doc, is_from_textbox=True
                    )

                elif local == "tbl":
                    # í…Œì´ë¸”ì´ë©´ í…Œì´ë¸” ì²˜ë¦¬
                    try:
                        self._handle_tables_enhanced(child, docx_obj, doc)
                    except Exception as e:
                        _log.debug(f"í…ìŠ¤íŠ¸ë°•ìŠ¤ ë‚´ í…Œì´ë¸” íŒŒì‹± ì‹¤íŒ¨: {e}")

                # (í•„ìš”ì‹œ ë‹¤ë¥¸ íƒœê·¸ë“¤: tbl, p ì™¸ì—ë„ shape/text ë“±)

        # ë¶€ëª¨ ë³µì›
        self.parents[level] = original_parent
        return

    def _handle_shape_content(
        self,
        element: BaseOxmlElement,
        docx_obj: DocxDocument,
        doc: DoclingDocument,
        owner_part=None,
    ) -> None:
        """Process shape content including tables, text, and images within shapes."""
        namespaces = {
            "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
            "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
            "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
            "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
            "mc": "http://schemas.openxmlformats.org/markup-compatibility/2006",
            "v": "urn:schemas-microsoft-com:vml",
            "wps": "http://schemas.microsoft.com/office/word/2010/wordprocessingShape",
            "w10": "urn:schemas-microsoft-com:office:word",
            "a14": "http://schemas.microsoft.com/office/drawing/2010/main",
        }

        # Look for text content in shapes
        shape_text_xpath = etree.XPath(
            ".//a:t|.//v:textbox//w:t|.//wps:txbx//w:t", namespaces=namespaces
        )
        text_elements = shape_text_xpath(element)
        
        # Look for tables in shapes
        shape_table_xpath = etree.XPath(
            ".//w:tbl|.//a:tbl", namespaces=namespaces
        )
        table_elements = shape_table_xpath(element)
        
        # Look for paragraphs in shapes
        shape_para_xpath = etree.XPath(
            ".//w:p|.//a:p", namespaces=namespaces
        )
        para_elements = shape_para_xpath(element)

        # Look for images in shapes (a:blip and VML v:imagedata)
        shape_img_xpath = etree.XPath(
            ".//a:blip|.//v:imagedata", namespaces=namespaces
        )
        image_elements = shape_img_xpath(element)

        if not text_elements and not table_elements and not para_elements and not image_elements:
            return

        # Extract all text content from the shape
        all_text_parts = []
        for text_elem in text_elements:
            if text_elem.text:
                all_text_parts.append(text_elem.text.strip())
        
        full_text = " ".join(all_text_parts).strip()
        
        # For longer text or complex content, create a shape group and process content
        level = self._get_level()
        shape_group = doc.add_group(
            label=GroupLabel.SECTION,
            parent=self.parents[level - 1],
            name="shape-content"
        )
        
        # Set this as the current parent temporarily
        original_parent = self.parents[level]
        self.parents[level] = shape_group

        # Process tables within the shape first
        for table_elem in table_elements:
            try:
                self._handle_tables_enhanced(table_elem, docx_obj, doc, owner_part=owner_part)
            except Exception as e:
                    _log.debug(f"Could not parse table in shape: {e}")

        # Process paragraphs within the shape
        for para_elem in para_elements:
            try:
                self._handle_text_elements(para_elem, docx_obj, doc)
            except Exception as e:
                    _log.debug(f"Could not parse paragraph in shape: {e}")

        # If we have text but no structured content, add it as plain text
        if full_text and not table_elements and not para_elements:
            # Check for duplicate content before adding
            if not self._is_duplicate_content(full_text):
                doc.add_text(
                    label=DocItemLabel.PARAGRAPH,
                    text=full_text,
                    parent=shape_group,
                    prov=ProvenanceItem(
                        page_no=1,
                        bbox=BoundingBox(l=0, t=0, r=1, b=1),
                        charspan=(0, len(full_text))
                    )
                )

        # Process images within the shape
        for img_elem in image_elements:
            try:
                local = etree.QName(img_elem).localname
                # Handle a:blip and v:imagedata uniformly by passing list to _handle_pictures
                blips = [img_elem] if local in ("blip", "imagedata") else []
                if blips:
                    self._handle_pictures(owner_part or docx_obj.part, docx_obj, blips, doc)
            except Exception as e:
                _log.debug(f"Could not parse image in shape: {e}")

        # Restore original parent
        self.parents[level] = original_parent
        return

    def _handle_equations_in_text(self, element, text):
        only_texts = []
        only_equations = []
        texts_and_equations = []
        for subt in element.iter():
            tag_name = etree.QName(subt).localname
            if tag_name == "t" and "math" not in subt.tag:
                if isinstance(subt.text, str):
                    only_texts.append(subt.text)
                    texts_and_equations.append(subt.text)
            elif "oMath" in subt.tag and "oMathPara" not in subt.tag:
                latex_equation = str(oMath2Latex(subt)).strip()
                if len(latex_equation) > 0:
                    only_equations.append(
                        self.equation_bookends.format(EQ=latex_equation)
                    )
                    texts_and_equations.append(
                        self.equation_bookends.format(EQ=latex_equation)
                    )

        if len(only_equations) < 1:
            return text, []

        if (
            re.sub(r"\s+", "", "".join(only_texts)).strip()
            != re.sub(r"\s+", "", text).strip()
        ):
            # If we are not able to reconstruct the initial raw text
            # do not try to parse equations and return the original
            return text, []

        # Insert equations into original text
        # This is done to preserve white space structure
        output_text = text[:]
        init_i = 0
        for i_substr, substr in enumerate(texts_and_equations):
            if len(substr) == 0:
                continue

            if substr in output_text[init_i:]:
                init_i += output_text[init_i:].find(substr) + len(substr)
            else:
                if i_substr > 0:
                    output_text = output_text[:init_i] + substr + output_text[init_i:]
                    init_i += len(substr)
                else:
                    output_text = substr + output_text

        return output_text, only_equations

    def _create_or_reuse_parent(
        self,
        *,
        doc: DoclingDocument,
        prev_parent: Optional[NodeItem],
        paragraph_elements: list,
    ) -> Optional[NodeItem]:
        return (
            doc.add_group(label=GroupLabel.INLINE, parent=prev_parent)
            if len(paragraph_elements) > 1
            else prev_parent
        )

    def _handle_text_elements(  # noqa: C901
        self,
        element: BaseOxmlElement,
        docx_obj: DocxDocument,
        doc: DoclingDocument,
        is_from_textbox: bool = False,
    ) -> None:
        paragraph = Paragraph(element, docx_obj)

        # Skip if from a textbox and this exact paragraph content was already processed
        raw_text = paragraph.text
        if is_from_textbox and raw_text:
            # Create a simple hash of content to detect duplicates
            content_hash = f"{len(raw_text)}:{raw_text[:50]}"
            if content_hash in self.processed_paragraph_content:
                return
            self.processed_paragraph_content.append(content_hash)

        text, equations = self._handle_equations_in_text(element=element, text=raw_text)

        if text is None:
            return
        paragraph_elements = self._get_paragraph_elements(paragraph)
        text = text.strip()

        # Common styles for bullet and numbered lists.
        # numbering.xml ê¸°ë°˜ìœ¼ë¡œ ìˆœì„œí˜• ì—¬ë¶€ íŒì •
        is_numbered = self._is_numbered_list(paragraph, docx_obj)
        p_style_id, p_level = self._get_label_and_level(paragraph)
        numid, ilevel = self._get_numId_and_ilvl(paragraph)

        if numid == 0:
            numid = None

        # Handle lists
        if (
            numid is not None
            and ilevel is not None
            and p_style_id not in ["Title", "Heading"]
        ):
            self._add_list_item(
                doc=doc,
                numid=numid,
                ilevel=ilevel,
                elements=paragraph_elements,
                is_numbered=is_numbered,
            )
            self._update_history(p_style_id, p_level, numid, ilevel)
            return
            
        
        elif (
            numid is None
            and self._prev_numid() is not None
            and p_style_id not in ["Title", "Heading"]
        ):  # Close list
            if self.level_at_new_list:
                for key in range(len(self.parents)):
                    if key >= self.level_at_new_list:
                        self.parents[key] = None
                self.level = self.level_at_new_list - 1
                self.level_at_new_list = None
            else:
                for key in range(len(self.parents)):
                    self.parents[key] = None
                self.level = 0

        if p_style_id in ["Title"]:
            for key in range(len(self.parents)):
                self.parents[key] = None
            self.parents[0] = doc.add_text(
                parent=None, label=DocItemLabel.TITLE, text=text, prov=ProvenanceItem(
                page_no=1,
                bbox=BoundingBox(l=0, t=0, r=1, b=1),
                charspan=(0, 0)
                )
            )
            self._update_history(p_style_id, p_level, numid, ilevel)
            return
            
        elif "Heading" in p_style_id:
            style_element = getattr(paragraph.style, "element", None)
            if style_element is not None:
                is_numbered_style = (
                    "<w:numPr>" in style_element.xml or "<w:numPr>" in element.xml
                )
            else:
                is_numbered_style = False
            # í—¤ë”© í…ìŠ¤íŠ¸ê°€ ë¹„ì–´ ìˆìœ¼ë©´ runs ê¸°ë°˜ìœ¼ë¡œ ì¬êµ¬ì„± ì‹œë„
            if len(text.strip()) == 0:
                reconstructed = " ".join([t for t, _, _ in paragraph_elements if isinstance(t, str) and len(t.strip()) > 0]).strip()
                if len(reconstructed) > 0:
                    text = reconstructed
            # ì—¬ì „íˆ ë¹„ì–´ ìˆìœ¼ë©´ ë¹ˆ ì œëª© ì¶”ê°€ë¥¼ ë°©ì§€
            if len(text.strip()) == 0:
                self._update_history(p_style_id, p_level, numid, ilevel)
                return
            # ë²ˆí˜¸ ìŠ¤íƒ€ì¼ì´ë©´ numbering.xml ê¸°ë°˜ ë¼ë²¨ì„ ê³„ì‚°í•´ì„œ ì›ë¬¸ í…ìŠ¤íŠ¸ ì•ì— ë¶™ì¸ë‹¤
            if is_numbered_style:
                num_for_label, ilvl_for_label = numid, ilevel
                if num_for_label is None or ilvl_for_label is None:
                    n2, i2 = self._get_style_numId_and_ilvl(paragraph)
                    num_for_label = num_for_label if num_for_label is not None else n2
                    ilvl_for_label = ilvl_for_label if ilvl_for_label is not None else i2
                label = self._build_number_label(num_for_label, ilvl_for_label, docx_obj)
                if label:
                    text = f"{label} {text}".strip()
            self._add_header(doc, p_level, text, is_numbered_style)

        elif len(equations) > 0:
            if (raw_text is None or len(raw_text.strip()) == 0) and len(text) > 0:
                # Standalone equation
                level = self._get_level()
                doc.add_text(
                    label=DocItemLabel.FORMULA,
                    parent=self.parents[level - 1],
                    text=text.replace("<eq>", "").replace("</eq>", ""),
                    prov=ProvenanceItem(
                    page_no=1,
                    bbox=BoundingBox(l=0, t=0, r=1, b=1),
                    charspan=(0, 0)
                    )
                )
                self._update_history(p_style_id, p_level, numid, ilevel)
                return
            else:
                # Inline equation
                level = self._get_level()
                inline_equation = doc.add_group(
                    label=GroupLabel.INLINE, parent=self.parents[level - 1]
                )
                text_tmp = text
                for eq in equations:
                    if len(text_tmp) == 0:
                        break

                    split_text_tmp = text_tmp.split(eq.strip(), maxsplit=1)

                    pre_eq_text = split_text_tmp[0]
                    text_tmp = "" if len(split_text_tmp) == 1 else split_text_tmp[1]

                    if len(pre_eq_text) > 0:
                        doc.add_text(
                            label=DocItemLabel.PARAGRAPH,
                            parent=inline_equation,
                            text=pre_eq_text,
                            prov=ProvenanceItem(
                            page_no=1,
                            bbox=BoundingBox(l=0, t=0, r=1, b=1),
                            charspan=(0, 0)
                            )
                        )
                    doc.add_text(
                        label=DocItemLabel.FORMULA,
                        parent=inline_equation,
                        text=eq.replace("<eq>", "").replace("</eq>", ""),
                        prov=ProvenanceItem(
                        page_no=1,
                        bbox=BoundingBox(l=0, t=0, r=1, b=1),
                        charspan=(0, 0)
                        )
                    )

                if len(text_tmp) > 0:
                    doc.add_text(
                        label=DocItemLabel.PARAGRAPH,
                        parent=inline_equation,
                        text=text_tmp.strip(),
                        prov=ProvenanceItem(
                        page_no=1,
                        bbox=BoundingBox(l=0, t=0, r=1, b=1),
                        charspan=(0, 0)
                        )
                    )
                self._update_history(p_style_id, p_level, numid, ilevel)
                return

        elif p_style_id in [
            "Paragraph",
            "Normal",
            "Subtitle",
            "Author",
            "DefaultText",
            "ListParagraph",
            "ListBullet",
            "Quote",
        ]:
            level = self._get_level()
            parent = self._create_or_reuse_parent(
                doc=doc,
                prev_parent=self.parents.get(level - 1),
                paragraph_elements=paragraph_elements,
            )
            for text, format, hyperlink in paragraph_elements:
                # Check for duplicate content before adding
                if not self._is_duplicate_content(text):
                    doc.add_text(
                        label=DocItemLabel.PARAGRAPH,
                        parent=parent,
                        text=text,
                        formatting=format,
                        hyperlink=hyperlink,
                        prov=ProvenanceItem(
                        page_no=1,
                        bbox=BoundingBox(l=0, t=0, r=1, b=1),
                        charspan=(0, 0)
                        )
                    )
            self._update_history(p_style_id, p_level, numid, ilevel)
            return

        else:
            # Text style names can, and will have, not only default values but user values too
            # hence we treat all other labels as pure text
            level = self._get_level()
            parent = self._create_or_reuse_parent(
                doc=doc,
                prev_parent=self.parents.get(level - 1),
                paragraph_elements=paragraph_elements,
            )
            for text, format, hyperlink in paragraph_elements:
                # Check for duplicate content before adding
                if not self._is_duplicate_content(text):
                    doc.add_text(
                        label=DocItemLabel.PARAGRAPH,
                        parent=parent,
                        text=text,
                        formatting=format,
                        hyperlink=hyperlink,
                        prov=ProvenanceItem(
                        page_no=1,
                        bbox=BoundingBox(l=0, t=0, r=1, b=1),
                        charspan=(0, 0)
                        )
                    )
            self._update_history(p_style_id, p_level, numid, ilevel)
            return

        self._update_history(p_style_id, p_level, numid, ilevel)
        return

    def _add_header(
        self,
        doc: DoclingDocument,
        curr_level: Optional[int],
        text: str,
        is_numbered_style: bool = False,
    ) -> None:
        level = self._get_level()
        if isinstance(curr_level, int):
            if curr_level > level:
                # add invisible group
                for i in range(level, curr_level):
                    self.parents[i] = doc.add_group(
                        parent=self.parents[i - 1],
                        label=GroupLabel.SECTION,
                        name=f"header-{i}",
                    )
            elif curr_level < level:
                # remove the tail
                for key in range(len(self.parents)):
                    if key >= curr_level:
                        self.parents[key] = None

            current_level = curr_level
            parent_level = curr_level - 1
            add_level = curr_level
        else:
            current_level = self.level
            parent_level = self.level - 1
            add_level = 1

        if is_numbered_style:
            if add_level in self.numbered_headers:
                self.numbered_headers[add_level] += 1
            else:
                self.numbered_headers[add_level] = 1
            # text = f"{self.numbered_headers[add_level]} {text}"

            # Reset deeper levels
            next_level = add_level + 1
            while next_level in self.numbered_headers:
                self.numbered_headers[next_level] = 0
                next_level += 1

            # Scan upper levels
            previous_level = add_level - 1
            while previous_level in self.numbered_headers:
                # MSWord convention: no empty sublevels
                # I.e., sub-sub section (2.0.1) without a sub-section (2.1)
                # is processed as 2.1.1
                if self.numbered_headers[previous_level] == 0:
                    self.numbered_headers[previous_level] += 1

                # text = f"{self.numbered_headers[previous_level]}.{text}"
                previous_level -= 1

        heading_prov = ProvenanceItem(
            page_no=1,
            bbox=BoundingBox(l=0, t=0, r=1, b=1),
            charspan=(0, 0)
        )
        self.parents[current_level] = doc.add_heading(
            parent=self.parents[parent_level],
            text=text,
            level=add_level,
            prov=heading_prov
        )
        return

    def _add_list_item(
        self,
        *,
        doc: DoclingDocument,
        numid: int,
        ilevel: int,
        elements: list,
        is_numbered: bool = False,
    ) -> None:
        enum_marker = ""

        level = self._get_level()
        prev_indent = self._prev_indent()
        if self._prev_numid() is None:  # Open new list
            self.level_at_new_list = level

            self.parents[level] = doc.add_group(
                label=GroupLabel.LIST, name="list", parent=self.parents[level - 1]
            )

            # Set marker and enumerated arguments if this is an enumeration element.
            # Set marker: numbering.xml ê¸°ë°˜ ë¼ë²¨ì„ ìš°ì„  ì‚¬ìš©
            # self.docx_obj ëŠ” __init__ì—ì„œ í•­ìƒ ì„¤ì •ë˜ë©° convert ì „ì— None ì•„ë‹˜
            assert self.docx_obj is not None
            label = self._build_number_label(numid, ilevel, self.docx_obj)
            if is_numbered and label:
                enum_marker = label
            else:
                # fallback: ë‹¨ìˆœ ì¦ê°€í˜•
                self.listIter += 1
                if is_numbered:
                    enum_marker = str(self.listIter) + "."
                    is_numbered = True
            new_parent = self._create_or_reuse_parent(
                doc=doc,
                prev_parent=self.parents[level],
                paragraph_elements=elements,
            )
            for text, format, hyperlink in elements:
                doc.add_list_item(
                    marker=enum_marker,
                    enumerated=is_numbered,
                    parent=new_parent,
                    text=text,
                    formatting=format,
                    hyperlink=hyperlink,
                    prov=ProvenanceItem(
                    page_no=1,
                    bbox=BoundingBox(l=0, t=0, r=1, b=1),
                    charspan=(0, 0)
                    )
                )

        elif (
            self._prev_numid() == numid
            and self.level_at_new_list is not None
            and prev_indent is not None
            and prev_indent < ilevel
        ):  # Open indented list
            for i in range(
                self.level_at_new_list + prev_indent + 1,
                self.level_at_new_list + ilevel + 1,
            ):
                # Determine if this is an unordered list or an ordered list.
                # Set GroupLabel.ORDERED_LIST when it fits.
                self.listIter = 0
                if is_numbered:
                    self.parents[i] = doc.add_group(
                        label=GroupLabel.ORDERED_LIST,
                        name="list",
                        parent=self.parents[i - 1],
                    )
                else:
                    self.parents[i] = doc.add_group(
                        label=GroupLabel.LIST, name="list", parent=self.parents[i - 1]
                    )

            # numbering.xml ê¸°ë°˜ ë¼ë²¨ ì¬ì‚¬ìš©
            assert self.docx_obj is not None
            label = self._build_number_label(numid, ilevel, self.docx_obj)
            if is_numbered and label:
                enum_marker = label
            else:
                self.listIter += 1
                if is_numbered:
                    enum_marker = str(self.listIter) + "."
                    is_numbered = True

            new_parent = self._create_or_reuse_parent(
                doc=doc,
                prev_parent=self.parents[self.level_at_new_list + ilevel],
                paragraph_elements=elements,
            )
            for text, format, hyperlink in elements:
                doc.add_list_item(
                    marker=enum_marker,
                    enumerated=is_numbered,
                    parent=new_parent,
                    text=text,
                    formatting=format,
                    hyperlink=hyperlink,
                    prov=ProvenanceItem(
                    page_no=1,
                    bbox=BoundingBox(l=0, t=0, r=1, b=1),
                    charspan=(0, 0)
                    )
                )
        elif (
            self._prev_numid() == numid
            and self.level_at_new_list is not None
            and prev_indent is not None
            and ilevel < prev_indent
        ):  # Close list
            for k, v in self.parents.items():
                if k > self.level_at_new_list + ilevel:
                    self.parents[k] = None

            # numbering.xml ê¸°ë°˜ ë¼ë²¨ ì¬ì‚¬ìš©
            assert self.docx_obj is not None
            label = self._build_number_label(numid, ilevel, self.docx_obj)
            if is_numbered and label:
                enum_marker = label
            else:
                self.listIter += 1
                if is_numbered:
                    enum_marker = str(self.listIter) + "."
                    is_numbered = True
            new_parent = self._create_or_reuse_parent(
                doc=doc,
                prev_parent=self.parents[self.level_at_new_list + ilevel],
                paragraph_elements=elements,
            )
            for text, format, hyperlink in elements:
                doc.add_list_item(
                    marker=enum_marker,
                    enumerated=is_numbered,
                    parent=new_parent,
                    text=text,
                    formatting=format,
                    hyperlink=hyperlink,
                    prov=ProvenanceItem(
                    page_no=1,
                    bbox=BoundingBox(l=0, t=0, r=1, b=1),
                    charspan=(0, 0)
                    )
                )
            self.listIter = 0

        elif self._prev_numid() == numid or prev_indent == ilevel:
            # numbering.xml ê¸°ë°˜ ë¼ë²¨ ì¬ì‚¬ìš©
            assert self.docx_obj is not None
            label = self._build_number_label(numid, ilevel, self.docx_obj)
            if is_numbered and label:
                enum_marker = label
            else:
                self.listIter += 1
                if is_numbered:
                    enum_marker = str(self.listIter) + "."
                    is_numbered = True
            new_parent = self._create_or_reuse_parent(
                doc=doc,
                prev_parent=self.parents[level - 1],
                paragraph_elements=elements,
            )
            for text, format, hyperlink in elements:
                # Add the list item to the parent group
                doc.add_list_item(
                    marker=enum_marker,
                    enumerated=is_numbered,
                    parent=new_parent,
                    text=text,
                    formatting=format,
                    hyperlink=hyperlink,
                    prov=ProvenanceItem(
                    page_no=1,
                    bbox=BoundingBox(l=0, t=0, r=1, b=1),
                    charspan=(0, 0)
                    )
                )
        return

    def _handle_pictures(
        self, owner_part, docx_obj: DocxDocument, drawing_blip: Any, doc: DoclingDocument
    ) -> None:
        def get_docx_image_info(drawing_blip: Any) -> tuple[Optional[bytes], Optional[str]]:
            """ì´ë¯¸ì§€ ë°ì´í„°ì™€ í˜•ì‹ ì •ë³´ë¥¼ ë°˜í™˜í•©ë‹ˆë‹¤."""
            image_data: Optional[bytes] = None
            image_format: Optional[str] = None
            
            rId = drawing_blip[0].get(
                "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
            )
            if rId in docx_obj.part.rels:
                # Access the image part using the relationship ID
                image_part = docx_obj.part.rels[rId].target_part
                image_data = image_part.blob  # Get the binary image data
                # Try to get content type to identify format
                image_format = getattr(image_part, 'content_type', None)
                
            return image_data, image_format
        def get_image_info_from_owner(owner_part, drawing_blip: Any) -> tuple[Optional[bytes], Optional[str]]:
            rId = drawing_blip[0].get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
            if not rId or rId not in owner_part.rels:
                return None, None
            image_part = owner_part.rels[rId].target_part
            image_data = image_part.blob
            image_format = getattr(image_part, 'content_type', None)
            return image_data, image_format

        def is_valid_image_format(image_format: Optional[str], image_data: Optional[bytes]) -> bool:
            """ì‹¤ì œ ì´ë¯¸ì§€ í˜•ì‹ì¸ì§€ í™•ì¸í•©ë‹ˆë‹¤."""
            if not image_format or not image_data:
                return False
                
            # XMLì´ë‚˜ ê¸°íƒ€ ë¹„ì´ë¯¸ì§€ í˜•ì‹ ì œì™¸
            non_image_formats = [
                'application/xml',
                'text/xml',
                'text/plain',
                'application/json',
                'text/html'
            ]
            
            if image_format.lower() in non_image_formats:
                return False
                
            # ë§¤ì§ ë„˜ë²„ë¡œ ì‹¤ì œ ì´ë¯¸ì§€ì¸ì§€ í™•ì¸
            if len(image_data) < 4:
                return False
                
            magic_bytes = image_data[:4]
            
            # XML ì‹œì‘ íŒ¨í„´ í™•ì¸
            if magic_bytes.startswith(b'<?xm') or magic_bytes.startswith(b'<xml'):
                return False
                
            # ì•Œë ¤ì§„ ì´ë¯¸ì§€ ë§¤ì§ ë„˜ë²„ë“¤
            image_signatures = [
                b'\x89PNG',           # PNG
                b'\xff\xd8\xff',      # JPEG
                b'GIF8',              # GIF
                b'BM',                # BMP
                b'RIFF',              # WebP (RIFF ì»¨í…Œì´ë„ˆ)
                b'\x00\x00\x01\x00', # ICO
                b'\xd7\xcd\xc6\x9a', # WMF
                b'\x01\x00\x00\x00', # EMF
                b'II*\x00',          # TIFF (little-endian)
                b'MM\x00*',          # TIFF (big-endian)
            ]
            
            # ë§¤ì§ ë„˜ë²„ ì¤‘ í•˜ë‚˜ë¼ë„ ì¼ì¹˜í•˜ë©´ ì´ë¯¸ì§€ë¡œ ê°„ì£¼
            for signature in image_signatures:
                if magic_bytes.startswith(signature):
                    return True
                    
            # content-typeì´ imageë¡œ ì‹œì‘í•˜ëŠ” ê²½ìš°ì—ë„ ì‹œë„
            if image_format.lower().startswith('image/'):
                return True
                
            return False

        level = self._get_level()
        # Open the BytesIO object with PIL to create an Image
        image_data, image_format = get_image_info_from_owner(owner_part, drawing_blip)
        
        # ì´ë¯¸ì§€ ë°ì´í„°ê°€ ì—†ê±°ë‚˜ í˜•ì‹ì´ Noneì¸ ê²½ìš°ì—ë„ add_picture í˜¸ì¶œ
        
        # if image_data is None:
        #     doc.add_picture(
        #         parent=self.parents[level - 1],
        #         caption=None,
        #         prov=ProvenanceItem(
        #         page_no=1,
        #         bbox=BoundingBox(l=0, t=0, r=1, b=1),
        #         charspan=(0, 0)
        #             )
        #     )
        #     return
            
        # ì‹¤ì œ ì´ë¯¸ì§€ì¸ì§€ í™•ì¸
        if not is_valid_image_format(image_format, image_data):
            # XMLì´ë‚˜ ê¸°íƒ€ ë©”íƒ€ë°ì´í„°ëŠ” ì´ë¯¸ì§€ë¡œ ì²˜ë¦¬í•˜ì§€ ì•ŠìŒ
            return
            
        try:
            assert image_data is not None
            image_bytes = BytesIO(image_data)
            image_bytes.seek(0)  # í¬ì¸í„°ë¥¼ ì‹œì‘ìœ¼ë¡œ ì´ë™
            pil_image = Image.open(image_bytes)
            doc.add_picture(
                parent=self.parents[level - 1],
                image=ImageRef.from_pil(image=pil_image, dpi=72),
                caption=None,
                prov=ProvenanceItem(
                page_no=1,
                bbox=BoundingBox(l=0, t=0, r=1, b=1),
                charspan=(0, 0)
                    )
            )
            
        except (UnidentifiedImageError, OSError) as e:
            print(f"Pillow failed to load image: {e}")
            print(f"Attempting Wand conversion for format: {image_format}")
            
            # WMF/EMF í˜•ì‹ ì²˜ë¦¬ ì‹œë„ (Wand ì‚¬ìš©)
            if WAND_AVAILABLE and image_format and ('wmf' in image_format.lower() or 'emf' in image_format.lower()):
                try:
                    with WandImage(blob=image_data) as wand_img:
                        # Convert to PNG format
                        wand_img.format = 'png'
                        png_blob = wand_img.make_blob()
                        
                        if png_blob:  # PNG ë°ì´í„°ê°€ ìˆëŠ”ì§€ í™•ì¸
                            png_bytes = BytesIO(png_blob)
                            png_bytes.seek(0)
                            pil_image = Image.open(png_bytes)
                            doc.add_picture(
                                parent=self.parents[level - 1],
                                image=ImageRef.from_pil(image=pil_image, dpi=72),
                                caption=None,
                                prov=ProvenanceItem(
                                page_no=1,
                                bbox=BoundingBox(l=0, t=0, r=1, b=1),
                                charspan=(0, 0)
                                )
                            )
                            return
                except (WandException, Exception) as wand_error:
                    print(f"Wand conversion failed: {wand_error}")
                
                # ë‹¤ë¥¸ í˜•ì‹ë„ Wandë¡œ ì‹œë„
                if WAND_AVAILABLE:
                    try:
                        with WandImage(blob=image_data) as wand_img:
                            # Convert to PNG format
                            wand_img.format = 'png'
                            png_blob = wand_img.make_blob()
                            
                            if png_blob:
                                png_bytes = BytesIO(png_blob)
                                png_bytes.seek(0)
                                pil_image = Image.open(png_bytes)
                                doc.add_picture(
                                    parent=self.parents[level - 1],
                                    image=ImageRef.from_pil(image=pil_image, dpi=72),
                                    caption=None,
                                    prov=ProvenanceItem(
                                    page_no=1,
                                    bbox=BoundingBox(l=0, t=0, r=1, b=1),
                                    charspan=(0, 0)
                                    )
                                )
                                return
                    except (WandException, Exception) as wand_error:
                        print(f"Wand fallback conversion failed: {wand_error}")
                
                # ìµœì¢…ì ìœ¼ë¡œ ë¹ˆ ì´ë¯¸ì§€ í”Œë ˆì´ìŠ¤í™€ë” ì¶”ê°€
                doc.add_picture(
                    parent=self.parents[level - 1],
                    caption=None,
                    prov=ProvenanceItem(
                    page_no=1,
                    bbox=BoundingBox(l=0, t=0, r=1, b=1),
                    charspan=(0, 0)
                        )
                    )
        return

