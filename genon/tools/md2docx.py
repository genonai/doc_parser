#!/usr/bin/env python3
"""마크다운 문서를 Word(.docx)로 변환한다.

gitbook_doc/*.md 처럼 GitHub-flavored Markdown 으로 쓰인 매뉴얼을 배포용 docx 로 뽑기 위한
스크립트다. 새 패키지 설치 없이 저장소 .venv 에 이미 있는 markdown + python-docx 만 쓴다.

사용법:
    ./.venv/bin/python genon/tools/md2docx.py <input.md> [-o output.docx] [--title "제목"]

파이프라인:
    md ──markdown(tables/fenced_code/sane_lists/attr_list)──▶ HTML ──html.parser──▶ python-docx

지원 요소: h1~h4 · 문단 · ul/ol(중첩) · 표 · 펜스 코드블록 · blockquote · hr ·
           인라인 code/strong/em/링크 · 이미지(상대경로 해석)
"""

from __future__ import annotations

import argparse
import datetime as _dt
import re
import sys
from html import unescape
from html.parser import HTMLParser
from pathlib import Path

try:
    import markdown
except ImportError:  # pragma: no cover - 실행 환경 안내
    sys.exit("markdown 패키지가 필요합니다. 저장소 .venv 의 python 으로 실행하세요:\n"
             "  ./.venv/bin/python genon/tools/md2docx.py ...")

try:
    import docx
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor
except ImportError:  # pragma: no cover
    sys.exit("python-docx 패키지가 필요합니다. 저장소 .venv 의 python 으로 실행하세요:\n"
             "  ./.venv/bin/python genon/tools/md2docx.py ...")


# ── 스타일 상수 ──────────────────────────────────────────────────────────
BODY_FONT = "맑은 고딕"
CODE_FONT = "D2Coding"          # 없으면 Word 가 대체 고정폭 폰트를 씀
CODE_FALLBACK_FONT = "Consolas"
CODE_SHADING = "F2F2F2"
LINK_COLOR = RGBColor(0x1A, 0x5F, 0xB4)
INLINE_CODE_COLOR = RGBColor(0xA0, 0x30, 0x30)
MAX_IMAGE_WIDTH_CM = 15.0
INDENT_PER_LEVEL_CM = 0.6

# 인라인 서식이 붙을 수 있는 태그
_INLINE_TAGS = {"strong", "b", "em", "i", "code", "a", "br", "img", "span", "del", "s"}
_BLOCK_TAGS = {"h1", "h2", "h3", "h4", "h5", "h6", "p", "pre", "ul", "ol", "li",
               "table", "thead", "tbody", "tr", "th", "td", "blockquote", "hr", "div"}


# ── HTML → 노드 트리 ─────────────────────────────────────────────────────
class _Node:
    __slots__ = ("tag", "attrs", "children", "text", "parent")

    def __init__(self, tag, attrs=None, parent=None):
        self.tag = tag
        self.attrs = dict(attrs or {})
        self.children: list[_Node] = []
        self.text = ""
        self.parent = parent

    def __repr__(self):  # pragma: no cover - 디버깅용
        return f"<{self.tag} {len(self.children)}>"


_VOID = {"br", "img", "hr"}


class _TreeBuilder(HTMLParser):
    """python-markdown 출력(잘 형성된 HTML)을 노드 트리로 만든다."""

    def __init__(self):
        super().__init__(convert_charrefs=True)
        self.root = _Node("root")
        self._cur = self.root

    def handle_starttag(self, tag, attrs):
        node = _Node(tag, dict(attrs), parent=self._cur)
        self._cur.children.append(node)
        if tag not in _VOID:
            self._cur = node

    def handle_startendtag(self, tag, attrs):
        self._cur.children.append(_Node(tag, dict(attrs), parent=self._cur))

    def handle_endtag(self, tag):
        if tag in _VOID:
            return
        node = self._cur
        while node is not self.root and node.tag != tag:
            node = node.parent
        if node is not self.root:
            self._cur = node.parent

    def handle_data(self, data):
        if not data:
            return
        node = _Node("#text", parent=self._cur)
        node.text = data
        self._cur.children.append(node)


def _raw_text(node: _Node) -> str:
    if node.tag == "#text":
        return node.text
    if node.tag == "br":
        return "\n"
    return "".join(_raw_text(c) for c in node.children)


# ── docx 저수준 헬퍼 ─────────────────────────────────────────────────────
def _set_shading(element, fill: str):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    element.append(shd)


def _set_paragraph_border(paragraph, color="D0D0D0", sz=4, sides=("top", "bottom", "left", "right")):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    for side in sides:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:space"), "4")
        el.set(qn("w:color"), color)
        borders.append(el)
    p_pr.append(borders)


def _add_hyperlink(paragraph, url: str, runs_spec: list[tuple[str, dict]]):
    """외부 링크 run 을 추가한다 (python-docx 에 기본 API 가 없어 직접 구성)."""
    part = paragraph.part
    r_id = part.relate_to(
        url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    for text, fmt in runs_spec:
        run = OxmlElement("w:r")
        r_pr = OxmlElement("w:rPr")
        color = OxmlElement("w:color")
        color.set(qn("w:val"), "1A5FB4")
        r_pr.append(color)
        underline = OxmlElement("w:u")
        underline.set(qn("w:val"), "single")
        r_pr.append(underline)
        if fmt.get("code"):
            rfonts = OxmlElement("w:rFonts")
            rfonts.set(qn("w:ascii"), CODE_FALLBACK_FONT)
            rfonts.set(qn("w:hAnsi"), CODE_FALLBACK_FONT)
            r_pr.append(rfonts)
        if fmt.get("bold"):
            r_pr.append(OxmlElement("w:b"))
        if fmt.get("italic"):
            r_pr.append(OxmlElement("w:i"))
        run.append(r_pr)
        t = OxmlElement("w:t")
        t.set(qn("xml:space"), "preserve")
        t.text = text
        run.append(t)
        hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _add_toc_field(paragraph):
    """Word 목차 필드(TOC). 문서를 열고 F9 를 누르면 채워진다."""
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "목차를 갱신하려면 이 영역을 선택하고 F9 를 누르세요."
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for el in (begin, instr, separate, placeholder, end):
        run._r.append(el)


# ── 인라인 렌더링 ────────────────────────────────────────────────────────
def _render_inline(paragraph, node: _Node, base_dir: Path, fmt: dict | None = None):
    fmt = dict(fmt or {})
    for child in node.children:
        tag = child.tag
        if tag == "#text":
            if not child.text:
                continue
            run = paragraph.add_run(child.text)
            _apply_fmt(run, fmt)
        elif tag == "br":
            paragraph.add_run().add_break()
        elif tag in ("strong", "b"):
            _render_inline(paragraph, child, base_dir, {**fmt, "bold": True})
        elif tag in ("em", "i"):
            _render_inline(paragraph, child, base_dir, {**fmt, "italic": True})
        elif tag in ("del", "s"):
            _render_inline(paragraph, child, base_dir, {**fmt, "strike": True})
        elif tag == "code":
            run = paragraph.add_run(_raw_text(child))
            _apply_fmt(run, {**fmt, "code": True})
        elif tag == "a":
            href = child.attrs.get("href", "")
            text = _raw_text(child)
            if href.startswith(("http://", "https://")):
                is_code = any(c.tag == "code" for c in child.children)
                _add_hyperlink(paragraph, href, [(text, {**fmt, "code": is_code})])
            elif href.startswith("#"):
                # 같은 문서 안의 앵커 — 종이 문서에서는 의미가 없으므로 텍스트만 남긴다.
                _render_inline(paragraph, child, base_dir, fmt)
            else:
                # 로컬 상대 링크(다른 md 파일 등)는 "텍스트 (경로)" 로 남긴다.
                _render_inline(paragraph, child, base_dir, fmt)
                if href and text != href:
                    run = paragraph.add_run(f" ({href})")
                    _apply_fmt(run, {**fmt, "small": True})
        elif tag == "img":
            _insert_image(paragraph, child, base_dir)
        else:
            _render_inline(paragraph, child, base_dir, fmt)


def _apply_fmt(run, fmt: dict):
    run.bold = bool(fmt.get("bold"))
    run.italic = bool(fmt.get("italic"))
    if fmt.get("strike"):
        run.font.strike = True
    if fmt.get("code"):
        run.font.name = CODE_FALLBACK_FONT
        run._element.rPr.rFonts.set(qn("w:eastAsia"), CODE_FALLBACK_FONT)
        run.font.size = Pt(9.5)
        run.font.color.rgb = INLINE_CODE_COLOR
    if fmt.get("small"):
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)


def _insert_image(paragraph, node: _Node, base_dir: Path):
    src = node.attrs.get("src", "")
    if src.startswith(("http://", "https://")):
        run = paragraph.add_run(f"[이미지: {src}]")
        _apply_fmt(run, {"small": True})
        return
    path = (base_dir / src).resolve()
    if not path.is_file():
        run = paragraph.add_run(f"[이미지 없음: {src}]")
        _apply_fmt(run, {"small": True})
        return
    try:
        paragraph.add_run().add_picture(str(path), width=Cm(MAX_IMAGE_WIDTH_CM))
    except Exception as exc:  # 손상된 이미지 등
        run = paragraph.add_run(f"[이미지 삽입 실패: {src} ({exc})]")
        _apply_fmt(run, {"small": True})


# ── 블록 렌더링 ──────────────────────────────────────────────────────────
class _Renderer:
    def __init__(self, document, base_dir: Path):
        self.doc = document
        self.base_dir = base_dir

    def render(self, node: _Node, indent: int = 0, quote: bool = False):
        for child in node.children:
            self.render_block(child, indent=indent, quote=quote)

    def render_block(self, node: _Node, indent: int = 0, quote: bool = False):
        tag = node.tag

        if tag == "#text":
            if node.text.strip():
                self._paragraph(node, indent, quote)
            return

        if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
            level = min(int(tag[1]), 4)
            heading = self.doc.add_heading(level=level)
            _render_inline(heading, node, self.base_dir)
            return

        if tag == "p":
            self._paragraph(node, indent, quote)
            return

        if tag == "pre":
            self._code_block(node, indent)
            return

        if tag in ("ul", "ol"):
            self._list(node, indent, quote, ordered=(tag == "ol"))
            return

        if tag == "table":
            self._table(node)
            return

        if tag == "blockquote":
            self.render(node, indent=indent, quote=True)
            return

        if tag == "hr":
            p = self.doc.add_paragraph()
            _set_paragraph_border(p, color="BBBBBB", sz=6, sides=("bottom",))
            return

        if tag in ("div", "root"):
            self.render(node, indent=indent, quote=quote)
            return

        # 알 수 없는 블록: 인라인으로 처리
        self._paragraph(node, indent, quote)

    # ── 개별 블록 ──
    def _paragraph(self, node: _Node, indent: int, quote: bool):
        if not _raw_text(node).strip() and not any(c.tag == "img" for c in node.children):
            return
        p = self.doc.add_paragraph()
        if node.tag == "#text":
            p.add_run(node.text.strip())
        else:
            _render_inline(p, node, self.base_dir)
        p.paragraph_format.left_indent = Cm(INDENT_PER_LEVEL_CM * indent + (0.5 if quote else 0))
        if quote:
            _set_paragraph_border(p, color="C8C8C8", sz=12, sides=("left",))
            for run in p.runs:
                if run.font.color.rgb is None:
                    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    def _code_block(self, node: _Node, indent: int):
        text = _raw_text(node).rstrip("\n")
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(INDENT_PER_LEVEL_CM * indent + 0.2)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run(text)
        run.font.name = CODE_FALLBACK_FONT
        run._element.rPr.rFonts.set(qn("w:eastAsia"), CODE_FALLBACK_FONT)
        run.font.size = Pt(8)          # ASCII 다이어그램이 접히지 않도록 작게
        p.paragraph_format.line_spacing = 1.0
        _set_shading(p._p.get_or_add_pPr(), CODE_SHADING)
        _set_paragraph_border(p, color="DDDDDD", sz=4)

    def _list(self, node: _Node, indent: int, quote: bool, ordered: bool):
        # 번호 목록은 Word 의 "List Number" 스타일을 쓰지 않는다 — 그 스타일은 문서 전체에서
        # 하나의 번호 시퀀스를 공유해, 뒤쪽 목록이 1 이 아니라 이전 목록에 이어진 번호로 시작한다.
        # 대신 "1. " 접두어를 직접 찍어 목록마다 항상 1 부터 시작하게 한다.
        style = None if ordered else ("List Bullet" if indent == 0
                                      else f"List Bullet {min(indent + 1, 3)}")
        start = _list_start(node)
        for idx, item in enumerate(c for c in node.children if c.tag == "li"):
            inline_children = _Node("li")
            nested_blocks: list[_Node] = []
            for child in item.children:
                if child.tag in _BLOCK_TAGS and child.tag != "p":
                    nested_blocks.append(child)
                elif child.tag == "p":
                    inline_children.children.extend(child.children)
                else:
                    inline_children.children.append(child)
            if style is None:
                p = self.doc.add_paragraph()
                marker = p.add_run(f"{start + idx}. ")
                marker.bold = True
            else:
                try:
                    p = self.doc.add_paragraph(style=style)
                except KeyError:
                    p = self.doc.add_paragraph(style="List Bullet")
            _render_inline(p, inline_children, self.base_dir)
            p.paragraph_format.left_indent = Cm(0.8 + INDENT_PER_LEVEL_CM * indent)
            if style is None:
                p.paragraph_format.first_line_indent = Cm(-0.55)
            for block in nested_blocks:
                self.render_block(block, indent=indent + 1, quote=quote)

    def _table(self, node: _Node):
        rows: list[list[_Node]] = []
        header_count = 0
        for section in _iter_desc(node, {"tr"}):
            cells = [c for c in section.children if c.tag in ("td", "th")]
            if not cells:
                continue
            rows.append(cells)
            if all(c.tag == "th" for c in cells):
                header_count = max(header_count, len(rows))
        if not rows:
            return
        ncols = max(len(r) for r in rows)
        table = self.doc.add_table(rows=len(rows), cols=ncols)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        # 기본은 열 폭 균등 분할이라 짧은 열은 남고 긴 열은 눌린다. 내용 기반 자동 맞춤으로 바꾼다.
        table.autofit = True
        tbl_pr = table._tbl.tblPr
        layout = OxmlElement("w:tblLayout")
        layout.set(qn("w:type"), "autofit")
        tbl_pr.append(layout)
        for r_idx, cells in enumerate(rows):
            for c_idx in range(ncols):
                cell = table.cell(r_idx, c_idx)
                cell.paragraphs[0].text = ""
                if c_idx >= len(cells):
                    continue
                src = cells[c_idx]
                p = cell.paragraphs[0]
                _render_inline(p, src, self.base_dir)
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(1)
                for run in p.runs:
                    if run.font.size is None:
                        run.font.size = Pt(9)
                if src.tag == "th":
                    _set_shading(cell._tc.get_or_add_tcPr(), "EDEDED")
                    for run in p.runs:
                        run.bold = True
        self.doc.add_paragraph()


def _list_start(node: _Node) -> int:
    """`<ol start="3">` 같은 시작 번호를 존중한다 (markdown 의 `3.` 시작 목록)."""
    try:
        return int(node.attrs.get("start", 1))
    except (TypeError, ValueError):
        return 1


def _iter_desc(node: _Node, tags: set[str]):
    for child in node.children:
        if child.tag in tags:
            yield child
        else:
            yield from _iter_desc(child, tags)


# ── 문서 조립 ────────────────────────────────────────────────────────────
def _setup_styles(document):
    normal = document.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    heading_sizes = {"Heading 1": 20, "Heading 2": 15, "Heading 3": 12.5, "Heading 4": 11}
    for name, size in heading_sizes.items():
        style = document.styles[name]
        style.font.name = BODY_FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
        style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
        style.font.size = Pt(size)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(14 if size >= 15 else 10)
        style.paragraph_format.space_after = Pt(6)


def _add_cover(document, title: str, subtitle: str):
    for _ in range(4):
        document.add_paragraph()
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.size = Pt(26)
    run.bold = True
    p2 = document.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(subtitle)
    run2.font.size = Pt(12)
    run2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    p3 = document.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run(_dt.date.today().isoformat())
    run3.font.size = Pt(10)
    run3.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    document.add_page_break()
    heading = document.add_paragraph()
    hrun = heading.add_run("목차")
    hrun.bold = True
    hrun.font.size = Pt(16)
    _add_toc_field(document.add_paragraph())
    document.add_page_break()


def _lift_fences_out_of_blockquotes(text: str) -> str:
    """인용문(`>`) 안의 펜스 코드블록을 인용문 밖으로 끌어낸다.

    GitHub 은 `> ```bash` 를 코드블록으로 렌더하지만 python-markdown 의 fenced_code 는
    blockquote 안에서 동작하지 않아 `<code>bash echo hi</code>` 같은 인라인 코드가 된다.
    (그대로 두면 docx 에서 코드블록이 통째로 뭉개진다.)
    원본 md 는 올바른 GFM 이므로 소스를 고치는 대신 변환 단계에서 처리한다 — 해당 펜스만
    인용문 밖의 독립 코드블록으로 분리하고, 앞뒤 인용문은 그대로 유지한다.
    """
    out: list[str] = []
    lines = text.split("\n")
    i = 0
    fence_re = re.compile(r"^(\s*)(```+|~~~+)(.*)$")
    while i < len(lines):
        line = lines[i]
        if not line.startswith(">"):
            out.append(line)
            i += 1
            continue
        # 인용문 블록 수집
        block: list[str] = []
        while i < len(lines) and (lines[i].startswith(">") or lines[i].strip() == ">"):
            block.append(lines[i])
            i += 1
        # 인용 마커 제거한 본문
        inner = [ln[1:].removeprefix(" ") if ln.startswith(">") else ln for ln in block]
        if not any(fence_re.match(ln) for ln in inner):
            out.extend(block)
            continue
        # 펜스를 경계로 인용/코드 구간을 나눈다
        buf_quote: list[str] = []
        j = 0
        while j < len(inner):
            m = fence_re.match(inner[j])
            if not m:
                buf_quote.append(inner[j])
                j += 1
                continue
            if buf_quote:
                out.extend(f"> {ln}".rstrip() for ln in buf_quote)
                out.append("")
                buf_quote = []
            indent, fence, info = m.group(1), m.group(2), m.group(3)
            out.append(f"{fence}{info}")
            j += 1
            while j < len(inner):
                closing = fence_re.match(inner[j])
                if closing and closing.group(2)[0] == fence[0] and not closing.group(3).strip():
                    j += 1
                    break
                out.append(inner[j].removeprefix(indent))
                j += 1
            out.append(fence)
            out.append("")
        if buf_quote:
            out.extend(f"> {ln}".rstrip() for ln in buf_quote)
    return "\n".join(out)


def convert(md_path: Path, out_path: Path, title: str | None = None,
            subtitle: str = "", cover: bool = True) -> Path:
    text = _lift_fences_out_of_blockquotes(md_path.read_text(encoding="utf-8"))
    html = markdown.markdown(
        text,
        extensions=["tables", "fenced_code", "sane_lists", "attr_list", "md_in_html"],
    )
    builder = _TreeBuilder()
    builder.feed(html)
    builder.close()

    document = Document()
    _setup_styles(document)

    tree = builder.root
    if title is None:
        first_h1 = next((n for n in tree.children if n.tag == "h1"), None)
        title = _raw_text(first_h1).strip() if first_h1 is not None else md_path.stem

    if cover:
        _add_cover(document, title, subtitle)
        # 표지에 이미 제목을 넣었으므로 본문 첫 h1 은 건너뛴다.
        first_h1 = next((n for n in tree.children if n.tag == "h1"), None)
        if first_h1 is not None:
            tree.children.remove(first_h1)

    _Renderer(document, md_path.parent).render(tree)
    document.save(str(out_path))
    return out_path


def main(argv=None) -> int:
    parser = argparse.ArgumentParser(
        description="마크다운(.md) 문서를 Word(.docx)로 변환한다.")
    parser.add_argument("input", type=Path, help="입력 .md 경로")
    parser.add_argument("-o", "--output", type=Path, default=None,
                        help="출력 .docx 경로 (기본: 입력과 같은 위치·같은 이름)")
    parser.add_argument("--title", default=None, help="표지 제목 (기본: 첫 h1)")
    parser.add_argument("--subtitle", default="", help="표지 부제")
    parser.add_argument("--no-cover", action="store_true", help="표지·목차 페이지 생략")
    args = parser.parse_args(argv)

    md_path: Path = args.input
    if not md_path.is_file():
        print(f"입력 파일을 찾을 수 없습니다: {md_path}", file=sys.stderr)
        return 1
    out_path: Path = args.output or md_path.with_suffix(".docx")
    out_path.parent.mkdir(parents=True, exist_ok=True)

    convert(md_path, out_path, title=args.title, subtitle=args.subtitle,
            cover=not args.no_cover)
    print(f"생성 완료: {out_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
